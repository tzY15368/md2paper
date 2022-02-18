from __future__ import annotations
from ast import Bytes
from io import BytesIO, StringIO
from typing import Union,List,Tuple
import docx
from docx.shared import Inches,Cm
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import lxml
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree
import latex2mathml.converter
from PIL import Image as PILImage
import logging
import os
import sys

SRC_ROOT = os.path.split(os.path.split(os.path.abspath(__file__))[0])[0]
logging.debug(f"resource root:{SRC_ROOT}")
def latex_to_word(latex_input):
    return etree.fromstring(latex_input)

class DocNotSetException(Exception):
    pass
class DocManager():
    __doc_target = None
    @classmethod
    # doc_target: path-like string, file-like object or docx.Document
    def set_doc(cls,doc_target:Union[docx.Document,str,BytesIO]):
        if type(doc_target)==str:
            actual_path = os.path.join(SRC_ROOT,doc_target)
            logging.info(f"reading from template:{actual_path}")
            cls.__doc_target = docx.Document(actual_path)
        elif type(doc_target) == docx.Document:
            cls.__doc_target = doc_target
        elif type(doc_target) == BytesIO:
            cls.__doc_target = docx.Document(doc_target)
        else:
            raise TypeError(f"invalid doc target: expecting str or docx.Document type,\
                 got {type(doc_target)}")
        cls.__clear_tables()

    @classmethod
    def get_doc(cls)->docx.Document:
        if not cls.__doc_target:
            raise DocNotSetException("doc target is not set, call DM.set_doc")
        return cls.__doc_target

    @classmethod
    def __clear_tables(cls):
        # delete all tables on startup as we don't need them
        for i in range(len(cls.get_doc().tables)):
            t = cls.get_doc().tables[0]._element
            t.getparent().remove(t)
            t._t = t._element = None

    @classmethod
    def delete_paragraph_by_index(cls, index):
        p = cls.get_doc().paragraphs[index]._element
        p.getparent().remove(p)
        p._p = p._element = None    

    @classmethod
    def get_anchor_position(cls,anchor_text:str,anchor_style_name="")->int:
        # FIXME: 需要优化
        # 目前被设计成无状态的，只依赖template.docx文件以便测试，增加了性能开销
        # USE-WITH-CARE
        # 只靠标题的anchor-text找paragraph很容易找错，用的时候注意
        i = -1
        for _i,paragraph in enumerate(cls.get_doc().paragraphs):
            if anchor_text in paragraph.text:
                if (not anchor_style_name) or (paragraph.style.name == anchor_style_name):
                    i = _i
                    break
                        
        if i==-1: raise ValueError(f"anchor `{anchor_text}` not found") 
        return i + 1

    @classmethod
    def get_paragraph(cls,offset:int):
        return cls.get_doc().paragraphs[offset]

    # https://stackoverflow.com/questions/51360649/how-to-update-table-of-contents-in-docx-file-with-python-on-linux?rq=1
    @classmethod
    def update_toc(cls):
        namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        # add child to doc.settings element
        element_updatefields = lxml.etree.SubElement(
            cls.get_doc().settings.element, f"{namespace}updateFields"
        )
        element_updatefields.set(f"{namespace}val", "true")

    @classmethod
    def save(cls,out:Union[str,StringIO]):
        cls.__doc_target.save(out)

DM = DocManager

class BaseContent():

    # 在指定offset 【向上】填充paragraph，返回填充后最后一段的offset+1
    def render_paragraph(offset:int)->int:
        raise NotImplementedError
class Component():
    def __init__(self) -> None:
        self.__internal_text = Block()
    
    def get_internal_text(self)->Block:
        return self.__internal_text

    def get_default_location(self)->Block:
        return self.get_internal_text()

    def add_text(self, text:Union[str,List[BaseContent]], location:Block=None)->Union[BaseContent,None]:
        if not location:
            location = self.get_default_location()

        if type(text)==str:
            txt = Text.read(text)
            location.add_content(content_list=txt)
            if len(txt) > 0:
                return txt[-1]
            else:
                return None

        elif type(text)==list and sum([1 if isinstance(i,BaseContent) else 0 for i in text])==len(text):
            location.add_content(content_list=text)
            if len(text) > 0:
                return text[-1]
            else:
                return None
        else:
            raise TypeError("expected text/List[BaseContent]")

    # anchor_text: 用于找到插入段落位置
    # incr_next: 用于在插入新内容后往后删老模板当前段内容，
    # 直到删除到incr_kw往前incr_next个paragraph
    # incr_kw：见上面incr_next
    def render_template(self, anchor_text:str,  incr_next:int, incr_kw, anchor_style_name="")->int:
        offset = DM.get_anchor_position(anchor_text=anchor_text,anchor_style_name=anchor_style_name)
        i = 0
        while not incr_kw in DM.get_doc().paragraphs[offset+incr_next].text\
                 and (offset+incr_next)!=(len(DM.get_doc().paragraphs)-1):
            DM.delete_paragraph_by_index(offset)
            i = i+1
        logging.debug("Component:deleted {} lines when rendering template".format(i))
        return self.__internal_text.render_template(offset)


class Run():
    Normal = 1
    Italics = 2
    Bold = 4
    Formula = 8
    Superscript = 16
    Subscript = 32
    def __init__(self,text:str,style:int=0,tabstop:bool=False) -> None:
        self.text = text
        self.bold = style & self.Bold != 0
        self.italics = style & self.Italics != 0
        self.formula = style & self.Formula != 0
        self.subscript = style & self.Subscript != 0
        self.superscript = style & self.Superscript != 0
        self.__tabstop = tabstop
    
    def render_run(self,run):
        if self.formula and self.text:
            word_math = latex_to_word(self.text)
            run._element.append(word_math)
        else:
            run.text = self.text
            run.bold = self.bold
            run.italic = self.italics
            run.font.subscript = self.subscript
            run.font.superscript = self.superscript

    @classmethod
    def get_tabstop(cls)->Run:
        return Run("",tabstop=True)

    def is_tabstop(self)->bool:
        return self.__tabstop
class Text(BaseContent):
    # 换行会被该类内部自动处理

    def __init__(self, raw_text:str="",style:int=Run.Normal) -> None:
        self.__runs:List[Run] = []
        if raw_text:
            self.__runs.append(Run(raw_text,style))

    def add_run(self, run:Run)->Text:
        self.__runs.append(run)
        return self

    def add_hfill(self)->Text:
        self.__runs.append(Run.get_tabstop())
        return self

    def render_paragraph(self, offset: int) -> int:
        new_offset = offset
        p = DM.get_doc().paragraphs[new_offset].insert_paragraph_before()
        for run in self.__runs:
            if not run.is_tabstop():
                run.render_run(p.add_run())
            else:
                # https://stackoverflow.com/questions/58656450/how-to-use-tabletop-by-python-docx
                sec = DM.get_doc().sections[0]
                margin_end = docx.shared.Inches(
                    sec.page_width.inches - (sec.left_margin.inches + sec.right_margin.inches))
                tab_stops = p.paragraph_format.tab_stops
                # adding new tab stop, to the end point, and making sure that it's `RIGHT` aligned.
                tab_stops.add_tab_stop(margin_end, docx.enum.text.WD_TAB_ALIGNMENT.RIGHT)
                
        p.paragraph_format.first_line_indent = Cm(0.82)
        new_offset = new_offset + 1
        return new_offset

    @classmethod
    def read(cls, txt:str)->List[Text]:
        return [Text(i) for i in txt.split('\n')]

class ImageData():
    def __init__(self,src:str,alt:str,width_ratio=0) -> None:
        # 如果提供了0-1之间的width ratio，则会覆盖dpi设定，
        # 宽度1则图片宽约等于可编辑区域宽度，不等于纸张宽度。
        self.img_src = src
        self.img_alt = alt

        self.dpi = 360
        self.MAX_WIDTH_INCHES = 6

        if not self.img_src:
            self.size = (0,0)
            self.size_inches = (0,0)
            logging.debug("empty image, alt={}".format(self.img_alt))
            return
            
        img = PILImage.open(self.img_src)
        self.size = img.size
        img.close()

        img_size_ratio = self.size[0]/self.size[1]
        if width_ratio < 0 or width_ratio > 1 :
            raise ValueError("invalid image width ratio, expecting range[0,1]")
        if width_ratio != 0:
            width_inches = self.MAX_WIDTH_INCHES*width_ratio
            height_inches = width_inches/img_size_ratio
            self.size_inches = (width_inches,height_inches)
        else:
            result = (self.size[0]/self.dpi,self.size[1]/self.dpi)
            if result[0] > self.MAX_WIDTH_INCHES:
                result = (self.MAX_WIDTH_INCHES,self.MAX_WIDTH_INCHES/img_size_ratio)
            self.size_inches = result
        logging.debug(f"image size:{self.size_inches[0]},{self.size_inches[1]}")

    # returns width,height in Inches
    def get_size_in_doc(self)->Tuple[Inches]:
        return map(Inches,self.size_inches)
class Image(BaseContent):
    def __init__(self,data:List[ImageData]) -> None:
        super().__init__()
        self.__images = data
    
    def render_paragraph(self, offset: int) -> int:
        new_offset = offset
        for img in self.__images:
            DM.get_doc().paragraphs[new_offset].insert_paragraph_before()
            new_offset = new_offset + 1
            
            p = DM.get_doc().paragraphs[new_offset].insert_paragraph_before()
            new_offset = new_offset + 1
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER 
            p.style = DM.get_doc().styles['图名中文']
            if img.img_src:
                r = p.add_run()
                r.add_picture(img.img_src,*img.get_size_in_doc())

                p = DM.get_doc().paragraphs[new_offset].insert_paragraph_before()
                new_offset = new_offset + 1
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER 
                p.style = DM.get_doc().styles['图名中文']
            
            p.add_run().add_text(img.img_alt)
        
        # 结尾再换
        DM.get_paragraph(new_offset).insert_paragraph_before()
        new_offset = new_offset + 1

        return new_offset

# Row of Table
class Row():
    def __init__(self,data:List[str],top_border:bool=False) -> None:
        self.row:List[str] = data
        self.has_top_border = top_border

class Formula(BaseContent):
    def __init__(self,title:str,formula:str) -> None:
        super().__init__()
        self.__title:str = title
        self.__formula:str = formula
        
    def render_paragraph(self,offset: int) -> int:
        new_offset = offset
        p = DM.get_doc().paragraphs[new_offset].insert_paragraph_before()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table = DM.get_doc().add_table(rows=1,cols=3)
        p._p.addnext(table._tbl)
        DM.delete_paragraph_by_index(new_offset)

        # 公式cell
        if self.__title:
            cell_formula = table.rows[0].cells[1]
            cell_formula.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            empty = {'color':Table.white}
            Table.set_cell_border(
                cell_formula,
                top=empty,
                bottom=empty,
                start=empty,
                end=empty
            )
            _p = cell_formula.paragraphs[0]
            _p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = _p.add_run()
            Run(self.__formula,Run.Formula).render_run(r)

        # 标号cell
        cell_idx = table.rows[0].cells[2]
        cell_idx.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_idx_p = cell_idx.paragraphs[0]
        cell_idx_p.text = self.__title
        cell_idx_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        return new_offset


class Table(BaseContent):
    # table的最后一行默认有下边框，剩下依靠row的top-border自行决定
    black = "#000000"
    white = "#ffffff"
    def __init__(self,title:str, table:List[Row]) -> None:
        super().__init__()
        self.__title = title
        self.__table:List[Row] = table
        if len(table) < 1:
            raise ValueError("invalid table content")
        self.__cols = len(self.__table[0].row)
        self.__rows = len(self.__table)
    
    def render_paragraph(self, offset: int) -> int:
        new_offset = offset
        # 先换一行
        DM.get_doc().paragraphs[new_offset].insert_paragraph_before()
        new_offset = new_offset + 1
        p1 = DM.get_doc().paragraphs[new_offset].insert_paragraph_before()
        new_offset = new_offset + 1
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER 
        p1.style = DM.get_doc().styles['图名中文']
        # 先换一行
        p1.add_run().add_text(self.__title)

        p2 = DM.get_doc().paragraphs[new_offset].insert_paragraph_before()
        table = DM.get_doc().add_table(rows = self.__rows, cols = self.__cols, style='Table Grid')
        # 将table挪到paragrpah里
        p2._p.addnext(table._tbl)
        # 挪完删掉paragraph
        DM.delete_paragraph_by_index(new_offset)
        
        #new_offset = new_offset + 1
        
        # 结尾再换
        p1 = DM.get_doc().paragraphs[new_offset].insert_paragraph_before()
        new_offset = new_offset + 1

        # 填充内容, 编辑表格样式
        for i,row in enumerate(self.__table):
            for j,cell_str in enumerate(row.row):
                cell = table.rows[i].cells[j]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                Table.set_cell_border(
                    cell,
                    top={"val":'single','color':self.white if not row.has_top_border else self.black},
                    bottom = {"val":'single', "color":self.white if i!=self.__rows-1 else self.black},
                    start={"color":self.white},
                    end={"color":self.white}
                )
                if cell_str == None:
                    if i == 0:
                        raise ValueError("invalid empty field in row 0")
                    else:
                        # 上一行同一列的cell
                        other_cell = table.rows[i-1].cells[j]
                        cell.merge(other_cell)
                        other_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                else:
                    p = cell.paragraphs[0]
                    p.text = cell_str
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.style = DM.get_doc().styles['图名中文']

        return new_offset

    # https://stackoverflow.com/questions/33069697/how-to-setup-cell-borders-with-python-docx
    @classmethod
    def set_cell_border(cls, cell, **kwargs):
        """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # check for tag existnace, if none found, then create one
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        # list over all available tags
        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)

                # check for tag existnace, if none found, then create one
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)

                # looks like order of attributes is important
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))

class References(Component): #参考文献
    def render_template(self) -> int:
        ANCHOR = "参 考 文 献"
        incr_next = 1
        incr_kw = "附录A"
        offset_start = DM.get_anchor_position(ANCHOR)
        offset_end = super().render_template(ANCHOR, incr_next, incr_kw) -incr_next+1
        _style = DM.get_doc().styles['参考文献正文']
        for i in range(offset_start,offset_end):
            DM.get_doc().paragraphs[i].style = _style
        return offset_end  

"""
每一章是一个chapter，
每个chapter内标题后可以直接跟block或section，
每个section内标题后可以直接跟block或subsection，
每个subsection内标题后可以跟block
"""
class Block(): #content
    # 每个block是多个image，formula，text的组合，内部有序
    heading_1 = 1
    heading_2 = 2
    heading_3 = 3
    heading_4 = 4

    def __init__(self) -> None:
        self.__title:str = None
        self.__content_list:List[Union[Text,Image,Table,Formula]] = []
        self.__sub_blocks:List[Block] = []
        self.__id:int = None

    def set_id(self, id:int):
        self.__id = id

    # 由level决定标题的样式（heading1，2，3）
    def set_title(self,title:str,level:int) -> Block:
        self.__title = title
        if level not in range(0,5):
            raise ValueError("invalid heading level")
        self.__level = level
        return self
    
    def add_sub_block(self,block:Block)->Block:
        self.__sub_blocks.append(block)
        return block

    def get_last_sub_block(self)->Block:
        if len(self.__sub_blocks)==0:
            raise ValueError("no available sub-blocks")
        return self.__sub_blocks[-1]

    def add_content(self,content:Union[Text,Image,Table,Formula]=None,
            content_list:Union[List[Text],List[Image],List[Table],List[Formula]]=[]) -> Block:
        if content:
            self.__content_list.append(content)
        for i in content_list:
            self.__content_list.append(i)
        logging.debug(f"added content")
        return self

    # render_template是基于render_block的api，增加了嵌套blocks的渲染 以支持递归生成章节/段落，
    # 同时增加了对段落标题和段落号的支持
    # 顺序：先title，再自己的content-list，再自己的sub-block
    def render_template(self, offset:int)->int:
        new_offset = offset

        # 如果是一级，给头上（标题前面）增加分页符
        if self.__title and self.__level == self.heading_1:
            p = DM.get_doc().paragraphs[new_offset].insert_paragraph_before()
            run = p.add_run()
            run.add_break(WD_BREAK.PAGE)
            new_offset = new_offset + 1

        if self.__title:
            logging.debug(f"block(level={self.__level}) title: {self.__title}")
            p_title = DM.get_doc().paragraphs[new_offset].insert_paragraph_before()
            p_title.style = DM.get_doc().styles['Heading '+str(self.__level)]
            p_title.add_run()
            title_idx = "" if not self.__id else str(self.__id) + "  "
            p_title.runs[0].text= title_idx + self.__title
            new_offset = new_offset + 1
        
        new_offset = self.render_block(new_offset)

        logging.debug(f"this block has {len(self.__sub_blocks)} sub-blocks")
        for i,block in enumerate(self.__sub_blocks):
            new_offset = block.render_template(new_offset) 

        return new_offset


    # render_block是最底层的api，只将自己的content-list加到已有文档给定位置
    # render_block takes the desired paragraph position's offset,
    # renders the block with native elements: text, image and formulas,
    # and returns the final paragraph's offset
    
    def render_block(self, offset:int)->int:
        if not self.__content_list:
            return offset
        new_offset = offset
        for i,content in enumerate(self.__content_list):
            new_offset = content.render_paragraph(new_offset)
            _media_types = [Image,Formula,Table]
            if i <len(self.__content_list)-1 and type(content) in _media_types and type(self.__content_list[i+1]) in _media_types:
                # 多媒体内容之间也只空一行
                DM.delete_paragraph_by_index(new_offset)
                new_offset = new_offset - 1

        return new_offset