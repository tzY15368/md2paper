from io import BytesIO, StringIO
import markdown
from bs4 import BeautifulSoup, Comment
import logging
import re
from functools import reduce
import os
import bibtexparser
from bibtexparser.bparser import BibTexParser
from typing import Dict, List, Union
import pypandoc
import docx

from md2paper.mdext import MDExt
import md2paper.dut_paper as word

debug = False


# 检查

def assert_warning(e: bool, s: str):
    if not e:
        logging.warning(s)
    return e


def assert_error(e: bool, s: str):
    if not e:
        logging.error(s)
        exit(-1)
    return e


def log_error(s: str):
    return assert_error(False, s)


def log_warning(s: str):
    return assert_warning(False, s)


# 处理文本

def rbk(text: str):  # remove_blank
    # 删除换行符
    text = text.replace("\n", " ")
    text = text.replace("\r", "")
    text = text.strip(' ')

    cn_char = u'[\u4e00-\u9fa5。，：《》、（）“”‘’\u00a0]'
    # 中文字符后空格
    should_replace_list = re.compile(
        cn_char + u' +').findall(text)
    # 中文字符前空格
    should_replace_list += re.compile(
        u' +' + cn_char).findall(text)
    # 删除空格
    for i in should_replace_list:
        if i == u' ':
            continue
        new_i = i.strip(" ")
        text = text.replace(i, new_i)
    text = text.replace("\u00a0", " ")  # 替换为普通空格
    return text


def raw_text(runs):
    strs = [i["text"] for i in runs]
    return reduce(lambda x, y: x+y, strs)


def assemble_ps(ps):
    strs = []
    for (_, runs) in ps:
        strs.append(raw_text(runs))
    return reduce(lambda x, y: x+"\n"+y, strs)


def ref_items_list_unfold(ref_items_list: list):
    unfold_ref_items = {}
    for ref_items in ref_items_list:
        for ali in ref_items:
            assert_warning(ali not in unfold_ref_items,
                           "任意类型的引用别名应该唯一: " + ali)
            unfold_ref_items[ali] = ref_items[ali]
    return unfold_ref_items


def re_space(s: str):
    return re.compile("^ *{} *".format(s))

# 数据类型


class RefItem:
    IMG = "img"
    TABLE = "Table"
    MATH = "math"
    LITER = "literature"

    def __init__(self, index, text: str, type: str):
        self.index = str(index)
        self.text = text
        self.type = type


# 每个论文模块

class PaperPart:
    def __init__(self):
        self.contents = []
        self.block: word.Component = None
        self.file_dir: str = ""

    def set_file_dir(self, file_dir: str):
        self.file_dir = file_dir

    # 获取内容

    def load_contents(self, soup: BeautifulSoup): pass

    def _get_content_until(self, cur, until, ollevel=4):
        conts = []
        head_counter = [0]
        while cur != until:
            if cur.name == None:
                cur = cur.next_sibling
                continue
            if cur.name[0] == "h":  # h1 h2 h3
                head_counter, pair = self._process_headline(head_counter,
                                                            cur.name, cur.text)
                conts.append(pair)
            elif cur.name == "p":
                conts += self._process_ps(cur)
            elif cur.name == "table":
                table_name = raw_text(conts[-1][1])
                conts = conts[:-1]
                conts.append(self._process_table(table_name, cur))
            elif cur.name == "ol":
                conts += self._process_ol(cur, ollevel)
            elif cur.name == "math":
                math_title = raw_text(conts[-1][1])
                conts = conts[:-1]
                conts.append(self._process_math(math_title, cur))
            else:
                log_error("这是啥？" + cur.prettify())
            cur = cur.next_sibling
        return conts

    def _get_content_from(self, cur, ollevel=4):
        return self._get_content_until(cur, None, ollevel)

    # 处理标签

    def _process_headline(self, head_counter: List[int], h_label: str, headline: str):
        level = int(h_label[1:])
        assert_warning(1 <= level and level <= len(head_counter)+1,
                       "标题层级应该递进" + headline)
        if level == len(head_counter) + 1:  # new sub part
            head_counter.append(1)
        elif 1 <= level and level <= len(head_counter):  # new part
            head_counter[level-1] += 1
            head_counter = head_counter[:level]
        else:
            log_error("错误的标题编号")

        index = str(head_counter[0])
        for i in range(1, len(head_counter)):
            index += "." + str(head_counter[i])

        headline = headline.strip()
        assert_warning(headline[:len(index)] == index,
                       "没有编号或者编号错误: {} {}".format(h_label, headline))
        assert_warning(headline[len(index)] == " " and
                       headline[len(index)+1] != " ",
                       "MD 中编号后应该有一个空格: {} {}".format(h_label, headline))
        headline = headline[:len(index)] + "  " + rbk(headline[len(index)+1:])

        return head_counter, (h_label, headline)

    def _process_ps(self, p, ollevel=4):
        ps = []
        data = []
        for i in p.children:
            if i.name == None:
                if not hasattr(i,"text"):
                    setattr(i,"text",str(i))
                if i.text == "\n":
                    continue
                data.append({"type": "text", "text": rbk(i.text)})
            elif i.name == "strong":
                assert_warning(len(i.contents) == 1, "只允许粗斜体，不允许复杂嵌套")
                if i.contents[0].name == "em":
                    data.append({"type": "strong-em", "text": rbk(i.text)})
                else:
                    data.append({"type": "strong", "text": rbk(i.text)})
            elif i.name == "em":
                data.append({"type": "em", "text": rbk(i.text)})
            elif i.name == "math-inline":
                data.append({"type": "math-inline", "text": i.text})
            elif i.name == "ref":
                data.append({"type": "ref", "text": rbk(i.text)})
            else:  # 需要分段
                if data:
                    ps.append(("p", data))
                    data = []
                if i.name == "br":  # 分段
                    pass
                elif i.name == "img":  # 图片
                    ps.append(self._process_img(i))
                elif i.name == "ol":
                    ps += self._process_ol(i, ollevel)
                else:
                    log_error("缺了什么？" + str(i))
        if data:
            ps.append(("p", data))
        return ps

    def _process_img(self, img):
        if img["src"] == "":
            img_path = ""
        else:
            img_path = os.path.join(self.file_dir, img["src"])
        ali, title, ratio = self._split_title(img["alt"])
        return ("img", {"alias": ali,
                        "title": title,
                        "ratio": ratio,
                        "src": img_path})

    def _process_table(self, title, table):
        data = []
        # 表头，有上实线
        data.append(word.Row([rbk(i.text) for i in table.find("thead").find_all("th")],
                             top_border=True))
        has_border = True  # 表身第一行有上实线
        for tr in table.find("tbody").find_all("tr"):
            row = [rbk(i.text) for i in tr.find_all("td")]  # get all text
            row = list(map(lambda x: None if x == '' else x,
                           row))  # replace '' with None
            if has_border:
                data.append(word.Row(row, top_border=True))
                has_border = False
            else:
                is_border = True
                for i in row:
                    if i == None:
                        is_border = False
                        break
                    for j in i:
                        if j != '-':
                            is_border = False
                            break
                if is_border:
                    has_border = True  # 自定义的实线，下一行数据有上实线
                else:
                    data.append(word.Row(row))

        ali, title, _ = self._split_title(title)
        return ("table", {"alias": ali,
                          "title": title,
                          "data": data})

    def _process_lis(self, li, level):
        if not hasattr(li.contents[0],"text"):
            setattr(li.contents[0],"text",str(li.contents[0]))
        if (li.contents[0].text == "\n"):  # <p>
            conts = self._get_content_from(li.contents[0], level+1)
        else:  # text
            conts = self._process_ps(li, level+1)
        conts[0] = ("fh" + str(level), conts[0][1])
        return conts

    def _process_ol(self, ol, level):
        assert_error(level <= 5, "层次至多两层")
        datas = [self._process_lis(i, level)
                 for i in ol.find_all("li", recursive=False)]
        # make index
        for i in range(len(datas)):
            li_data = datas[i][0]
            if level == 4:
                li_data[1].insert(
                    0, {"type": "text", "text": "（{}） ".format(i+1)})
            else:
                assert_warning(i < 20, "层次二不能超过 20 项")
                li_data[1].insert(
                    0, {"type": "text", "text": "{} ".format(chr(i+0x2460))})  # get ①②..⑳
        data = reduce(lambda x, y: x + y, datas)
        return data

    def _process_math(self, title, math):
        return ("math", {"alias": title,
                         "title": "",
                         "text": math.text})

    def _split_title(self, title: str):
        sp = title.split(':')
        sp = [sp[0]] + sp[1].split(';')
        if len(sp) == 2:
            ali = sp[0]
            title = rbk(sp[1])
            ratio = 0
        elif len(sp) == 3:
            ali = sp[0]
            title = rbk(sp[1])
            ratio_s = rbk(sp[2])
            if ratio_s != "":
                ratio = int(ratio_s[:-1])
            else:
                ratio = 0
            assert_warning(0 <= ratio <= 100,
                           "图片占页面宽度应该在 [0%, 100%] 间: " + title)
        else:
            log_error("图、表的标题格式错误: " + title)

        return ali, title, ratio/100

    # 处理

    def check(self): pass

    def _math_pandoc_word(self):
        tmp_doc = "**tmp**.docx"

        # get math
        math_list: List[str] = []
        for name, cont in self.contents:
            if name == "p":
                for run in cont:
                    if run["type"] == "math-inline" and run["text"].strip() != "":
                        math_list.append(run["text"])
            elif name == "math" and cont["text"].strip() != "":
                math_list.append(cont["text"])

        # get word
        if math_list == []:
            return
        md_list = ["${}$".format(i.strip()) for i in math_list]
        md = reduce(lambda x, y: x+'\n\n'+y, md_list)
        pypandoc.convert_text(md, "docx", "md", outputfile=tmp_doc)
        doc = docx.Document(tmp_doc)
        paras_xml = [str(i._element.xml) for i in doc.paragraphs]
        os.remove(tmp_doc)

        oMath_head = "<m:oMath>"
        oMath_tail = "</m:oMath>"
        word_maths = [para_xml[para_xml.find(oMath_head):
                               para_xml.find(oMath_tail) + len(oMath_tail)]
                      for para_xml in paras_xml]
        word_maths_m: List[str] = []
        for word_math in word_maths:
            pos = word_math.find('>')
            word_math = word_math[:pos] + \
                ' xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"' + \
                word_math[pos:]
            word_maths_m.append(word_math)

        # put back
        count = 0
        for name, cont in self.contents:
            if name == "p":
                for run in cont:
                    if run["type"] == "math-inline" and run["text"].strip() != "":
                        run["text"] = word_maths_m[count]
                        count += 1
            elif name == "math" and cont["text"].strip() != "":
                cont["text"] = word_maths_m[count]
                count += 1

    def compile(self):
        self._math_pandoc_word()

    def _get_ref_items(self, conts, index_prefix: str = "") -> Dict[str, RefItem]:
        def get_index(index_prefix: str, chapter_cnt: int, item_cnt: int):
            if index_prefix == "":
                return "{}.{}".format(chapter_cnt, item_cnt)
            else:
                return "{}{}".format(index_prefix, item_cnt)

        def img_index() -> str:
            return get_index(index_prefix, chapter_cnt, img_cnt)

        def table_index() -> str:
            return get_index(index_prefix, chapter_cnt, img_cnt)

        def math_index() -> str:
            return get_index(index_prefix, chapter_cnt, img_cnt)

        ref_items = {}
        chapter_cnt = 0

        for name, cont in conts:
            if name == "h1":
                chapter_cnt += 1
                img_cnt = 0
                table_cnt = 0
                formula_cnt = 0
            elif name in ["img", "table", "math"]:
                ali = cont['alias']
                assert_warning(ali not in ref_items, "有重复别名" + ali)
                if name == "img":
                    img_cnt += 1
                    ref_items[ali] = RefItem(
                        img_index(), "图" + img_index(), RefItem.IMG)
                    cont['title'] = "图{}  {}".format(
                        img_index(), cont['title'])
                elif name == "table":
                    table_cnt += 1
                    ref_items[ali] = RefItem(
                        table_index(), "表" + table_index(), RefItem.TABLE)
                    cont['title'] = "表{}  {}".format(
                        table_index(), cont['title'])
                elif name == "math":
                    formula_cnt += 1
                    ref_items[ali] = RefItem(
                        math_index(), "式" + math_index(), RefItem.MATH)
                    cont['title'] = "（{}）".format(
                        table_index())
        return ref_items

    def get_ref_items(self):
        return self._get_ref_items(self.contents)

    def link_ref(self, ref_items: Dict[str, RefItem], liter_cnt: int) -> int:
        for name, cont in self.contents:
            if name not in ["p", "fh4", "fh5"]:
                continue
            is_text = False
            for run in cont:
                if run["type"] != "ref":
                    if run["type"] == "text" and run["text"].endswith("文献"):
                        is_text = True
                    continue
                ali = run["text"]
                if ali.find(",") == -1:
                    if ali not in ref_items:
                        ref_items[ali] = RefItem(
                            liter_cnt+1, "", RefItem.LITER)
                        liter_cnt += 1
                    ref_item = ref_items[ali]
                    if ref_item.type == RefItem.LITER:
                        run["text"] = "[{}]".format(
                            ref_item.index)
                        if is_text:
                            run["type"] = "text"
                        else:
                            run["type"] = "ref"
                    else:
                        run["type"] = "text"
                        run["text"] = ref_item.text
                else:
                    alis = ali.split(",")
                    for ali in alis:
                        if ali not in ref_items:
                            ref_items[ali] = RefItem(
                                liter_cnt+1, "", RefItem.LITER)
                            liter_cnt += 1
                    index_list = []
                    for ali in alis:
                        assert_error(ref_items[ali].type == RefItem.LITER,
                                     "只有参考文献可以一次引用多个: "+str(alis))
                        index_list.append(int(ref_items[ali].index))
                    index_list.sort()
                    index_list = [[x, x] for x in index_list]
                    sort_list = index_list[:1]
                    for index_pair in index_list[1:]:
                        if sort_list[-1][1]+1 == index_pair[0]:
                            sort_list[-1][1] = index_pair[1]
                        else:
                            sort_list.append(index_pair)
                    sort_list = [str(x[0]) if x[0] == x[1]
                                 else "{}-{}".format(x[0], x[1])
                                 for x in sort_list]
                    run["text"] = "[{}]".format(
                        reduce(lambda x, y: x+","+y, sort_list))
                    if is_text:
                        run["type"] = "text"
                    else:
                        run["type"] = "ref"
                is_text = False
        return liter_cnt

    # 渲染

    def _block_load_body(self, conts=None):
        if conts == None:
            conts = self.contents
        for (name, cont) in conts:
            if name == "h1":
                self.block.add_chapter(cont)
            elif name == "h2":
                self.block.add_section(cont)
            elif name == "h3":
                self.block.add_subsection(cont)
            elif name in ["p", "fh4", "fh5"]:
                if not debug:
                    para = word.Text()
                else:
                    para = word.Text(name)
                for run in cont:
                    if run["type"] == "text":
                        para.add_run(word.Run(run["text"], word.Run.Normal))
                    elif run["type"] == "strong":
                        para.add_run(word.Run(run["text"], word.Run.Bold))
                    elif run["type"] == "em":
                        para.add_run(word.Run(run["text"], word.Run.Italics))
                    elif run["type"] == "strong-em":
                        para.add_run(word.Run(run["text"],
                                              word.Run.Italics | word.Run.Bold))
                    elif run["type"] == "math-inline":
                        para.add_run(word.Run(run["text"], word.Run.Formula))
                    elif run["type"] == "ref":
                        para.add_run(
                            word.Run(run["text"], word.Run.Superscript))
                    else:
                        print("还没实现now", name)
                self.block.add_text([para])
            elif name == "img":
                img = word.Image(
                    [word.ImageData(cont["src"], cont["title"], cont["ratio"])])
                self.block.add_text([img])
            elif name == "table":
                table = word.Table(cont['title'], cont['data'])
                self.block.add_text([table])
            elif name == "math":
                formula = word.Formula(cont['title'], cont['text'])
                self.block.add_text([formula])
            else:
                print("还没实现now", name)

    def _block_load_contents(self):
        self._block_load_body()

    def render(self):
        self._block_load_contents()
        self.block.render_template()


class Paper:
    def __init__(self):
        self.parts: list[PaperPart] = []
        self.ref_items: Dict[str, Dict[str, str]] = {}
        self.file_dir: str = ""

    def load_md(self, md_path: str):
        with open(md_path, "r") as f:
            md_file = f.read()
        self.file_dir = os.path.dirname(md_path)
        for part in self.parts:
            part.set_file_dir(self.file_dir)
        md_html = markdown.markdown(md_file,
                                    tab_length=3,
                                    extensions=['markdown.extensions.tables',
                                                MDExt()])
        self.soup = BeautifulSoup(md_html, 'html.parser')
        for i in self.soup(text=lambda text: isinstance(text, Comment)):
            i.extract()  # 删除 html 注释

        if debug:
            with open("out.html", "w") as f:
                f.write(self.soup.prettify())

    def load_contents(self):
        for part in self.parts:
            part.load_contents(self.soup)

    def compile(self):
        for part in self.parts:
            part.compile()

    def render(self, doc: Union[str, BytesIO], out: Union[str, StringIO]):
        word.DM.set_doc(doc)

        for part in self.parts:
            part.render()

        word.DM.update_toc()
        word.DM.save(out)


'''
("h1", "something")
("h2", "something")
("h3", "something")

("fh4", like p)
("fh5", like p)

("p", [("text",      "something"),
       ("strong",    "something"),
       ("strong-em", "something"),
       ("em",        "something")])
("img",     (title, src))
("table",   (title, [Row]))
("formula", (title, "somthing"))
'''
