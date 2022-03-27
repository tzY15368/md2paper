import logging
from md2paper.md2paper import *

from typing import Dict
import datetime
from docx.enum.text import WD_LINE_SPACING


class Metadata(Component):
    school: str = None
    major: str = None
    name: str = None
    number: str = None
    teacher: str = None
    auditor: str = None
    __finish_date: str = None
    title_zh_CN: str = None
    title_en: str = None

    # 封面个人信息留空长度
    BLANK_LENGTH = 23

    def get_line_mapping(self) -> Dict[str, str]:
        data = {
            "学 院（系）：": self.school,
            "专       业：": self.major,
            "学 生 姓 名：": self.name,
            "学       号：": self.number,
            "指 导 教 师：": self.teacher,
            "评 阅 教 师：": self.auditor,
            "完 成 日 期：": self.finish_date
        }
        return data

    def get_title_mapping(self) -> Dict[str, str]:
        data = {
            "大连理工大学本科毕业设计（论文）题目": {
                "text": self.title_zh_CN,
                "max_len": 38
            },
            "The Subject of Undergraduate Graduation Project (Thesis) of DUT": {
                "text": self.title_en,
                "max_len": 66
            }
        }
        return data

    @property
    def finish_date(self):
        if self.__finish_date:
            return self.__finish_date
        else:
            now = datetime.datetime.now()
            return now.strftime("%Y年%m月%d日")

    @finish_date.setter
    def finish_date(self, value):
        self.__finish_date = value

    def __get_data_len(self, data: str) -> int:
        # 判断是否中文, 一个中文算两个字符
        def is_zh_CN(char) -> bool:
            return '\u4e00' <= char <= '\u9fa5'
        len = 0
        prev_char_type = False
        prev_char = None
        for char in data:
            current_char_type = is_zh_CN(char)
            len += 1
            if current_char_type:
                len += 1
            if current_char_type != prev_char_type and prev_char != None:
                len += 0.5
            prev_char_type = current_char_type
            prev_char = char
        return int(len)

    def __fill_blank(self, blank_length: int, data: str) -> str:
        """
        填充诸如 "学 生 姓 名：______________"的域
        **一个中文算两个字符
        **中文和数字之间会多一个空格
        """
        head_length = int((blank_length - self.__get_data_len(data)) / 2)
        if head_length < 0:
            raise ValueError("值过长")
        content = " " * head_length + data + " " * \
            (blank_length-self.__get_data_len(data)-head_length)
        return content

    def render_template(self):
        # 首先设置header
        for section in DM.get_doc().sections:
            p = section.header.paragraphs[0]
            if len(p.runs) == 0:
                continue
            p.runs[0].text = f"\t{self.title_zh_CN}\t"
            # 清空剩下的run
            for run in p.runs[1:]:
                run.text = ""

        mapping = self.get_title_mapping()
        for field in mapping:
            text = mapping[field].get('text')
            if not text:
                continue
            offset = DM.get_anchor_position(field) - 1
            DM.get_doc().paragraphs[offset].runs[0].text = text

            # 这里如果标题太长导致折行，则额外删去一行，以防止封面溢出到第二页
            logging.debug(
                f"metadata:text len = {self.__get_data_len(text)}, max len ={mapping[field]['max_len']}")
            if self.__get_data_len(text) >= mapping[field]['max_len']:
                DM.delete_paragraph_by_index(offset + 5)

        mapping = self.get_line_mapping()
        for field in mapping:
            if not mapping[field]:
                continue
            offset = DM.get_anchor_position(field) - 1
            data = self.__fill_blank(self.BLANK_LENGTH, mapping[field])
            DM.get_doc().paragraphs[offset].runs[-1].text = data


class Abstract(Component):

    def __init__(self) -> None:
        self.__text_en: Block = Block()
        self.__text_zh_CN: Block = Block()
        self.__keyword_en: str = None
        self.__keyword_zh_CN: str = None
        self.__title_zh_CN = ""
        self.__title_en = ""

    # 中文title被忽略，因为模板中没用上
    def set_title(self, zh_CN: str, en: str):
        self.__title_zh_CN = zh_CN
        self.__title_en = en

    def set_keyword(self, zh_CN: List[str], en: List[str]):
        SEPARATOR = "；"
        self.__keyword_en = SEPARATOR.join(en)
        self.__keyword_zh_CN = SEPARATOR.join(zh_CN)

    def add_text(self, zh_CN: str, en: str):
        self.__text_en.add_content(content_list=Text.read(en))
        self.__text_zh_CN.add_content(content_list=Text.read(zh_CN))

    def render_template(self) -> int:
        # 64开始是摘要正文
        #abs_cn_start = 64
        #abs_cn_end = self.__text_zh_CN.render_block(abs_cn_start)

        offset = DM.get_anchor_position(
            "摘    要", anchor_style_name="Heading 1")
        offset = self.__text_zh_CN.render_block(offset)

        while not DM.get_paragraph(offset+1).text.startswith("关键词："):
            DM.delete_paragraph_by_index(offset)

        # cn kw
        offset = offset + 1
        DM.get_paragraph(offset).runs[1].text = self.__keyword_zh_CN

        # en start
        offset = offset+4
        DM.get_paragraph(offset).runs[0].text = self.__title_en

        offset = offset + 3
        offset = self.__text_en.render_block(offset)

        # https://stackoverflow.com/questions/61335992/how-can-i-use-python-to-delete-certain-paragraphs-in-docx-document
        while not DM.get_paragraph(offset+1).text.startswith("Key Words："):
            DM.delete_paragraph_by_index(offset)

        # en kw
        offset = offset + 1
        # https://github.com/python-openxml/python-docx/issues/740
        delete_num = len(DM.get_doc().paragraphs[offset].runs) - 4
        for run in reversed(list(DM.get_doc().paragraphs[offset].runs)):
            DM.get_paragraph(offset)._p.remove(run._r)
            delete_num -= 1
            if delete_num < 1:
                break

        DM.get_paragraph(offset).runs[3].text = self.__keyword_en
        return offset+1


class Introduction(Component):
    def render_template(self) -> int:
        anchor_text = "引    言"
        incr_next = 2
        incr_kw = "正文格式说明"
        anchor_style_name = "Heading 1"
        return super().render_template(anchor_text, incr_next, incr_kw, anchor_style_name=anchor_style_name)


class MainContent(Component):  # 正文

    def __init__(self) -> None:
        super().__init__()
        self.__prev: Dict[str, Block] = {
            'blk': None,
            'chapter': None,
            'section': None,
            'subsection': None
        }

    def get_last_subblock(self) -> Block:
        if 'blk' not in self.__prev:
            raise ValueError("last blk is not yet set")
        return self.__prev['blk']

    # get_last_subblock
    def get_default_location(self) -> Block:
        if 'blk' not in self.__prev:
            raise ValueError("last blk is not yet set")
        return self.__prev['blk']

    # add_chapter returns the added chapter
    def add_chapter(self, title: str) -> Block:
        new_chapter = Block()

        # set prev pointers
        self.__prev['blk'] = new_chapter
        self.__prev['chapter'] = new_chapter

        new_chapter.set_title(title, Block.heading_1)
        return self.get_internal_text().add_sub_block(new_chapter)

    # add_section returns the added section
    def add_section(self, title: str, chapter: Block = None) -> Block:
        if not chapter:
            if self.__prev.get('chapter'):
                chapter = self.__prev['chapter']
            else:
                raise KeyError("chapter must be initialized before section")

        new_section = Block()

        # set prev pointers
        self.__prev['blk'] = new_section
        self.__prev['section'] = new_section

        new_section.set_title(title, Block.heading_2)
        return chapter.add_sub_block(new_section)

    # add_subsection returns the added subsection
    def add_subsection(self, title: str, section: Block = None) -> Block:
        if not section:
            if 'section' in self.__prev:
                section = self.__prev['section']
            else:
                raise KeyError("section must be initialized before subsection")

        new_subsection = Block()

        # set prev pointers
        self.__prev['blk'] = new_subsection
        self.__prev['subsection'] = new_subsection

        new_subsection.set_title(title, Block.heading_3)
        return section.add_sub_block(new_subsection)

    # 由于无法定位正文，需要先生成引言，再用引言返回的offset
    def render_template(self) -> int:
        anchor_text = "1  正文格式说明"
        incr_next = 3
        incr_kw = "结    论（设计类为设计总结"
        # 此处没有覆盖原有内容，因此还需要删去原有的大标题 1 正文格式……
        offset = super().render_template(anchor_text, incr_next, incr_kw)

        line_delete_count = 1
        pos = DM.get_anchor_position(anchor_text) - 1
        for i in range(line_delete_count):
            DM.delete_paragraph_by_index(pos)
        return offset - line_delete_count


class Conclusion(Component):
    def render_template(self, override_title: str = None) -> int:
        ANCHOR = "结    论（设计类为设计总结）"
        incr_next = 3
        incr_kw = "参 考 文 献"
        new_offset = super().render_template(ANCHOR, incr_next, incr_kw)
        if override_title:
            title_offset = DM.get_anchor_position(ANCHOR) - 1
            DM.get_doc().paragraphs[title_offset].runs[1].text = override_title
        return new_offset


class Appendixes(Component):  # 附录abcdefg, 是一种特殊的正文
    def __init__(self) -> None:
        super().__init__()

    def get_default_location(self) -> Block:
        return self.get_internal_text().get_last_sub_block()

    def add_appendix(self, appendix_title: str) -> Block:
        new_appendix = Block()
        new_appendix.set_title(appendix_title, level=Block.heading_1)
        self.get_internal_text().add_sub_block(new_appendix)
        return new_appendix

    def render_template(self) -> int:
        anchor_text = "附录A"
        anchor_style_name = "Heading 1"
        incr_next = 1
        incr_kw = "修改记录"
        offset = super().render_template(anchor_text, incr_next, incr_kw, anchor_style_name)
        # 此处没有覆盖原有内容，因此还需要删去原有的附录a那一页的3段

        line_delete_count = 1
        pos = DM.get_anchor_position(anchor_text=anchor_text)-1
        for i in range(line_delete_count):
            DM.delete_paragraph_by_index(pos)
        return offset - line_delete_count


class ChangeRecord(Component):  # 修改记录
    def render_template(self) -> int:
        ANCHOR = "修改记录"
        ANCHOR_STYLE = "Heading 1"
        incr_next = 0
        incr_kw = "致    谢"
        return super().render_template(ANCHOR, incr_next, incr_kw, anchor_style_name=ANCHOR_STYLE)


class Acknowledgments(Component):  # 致谢
    def render_template(self) -> int:
        ANCHOR = "致    谢"
        incr_next = 0

        # hack: 致谢已经到论文末尾，因此用无法匹配上的字符串直接让他删到最后一行
        incr_kw = "/\,.;'"
        return super().render_template(ANCHOR, incr_next, incr_kw)


class References(Component):  # 参考文献
    def render_template(self) -> int:
        ANCHOR = "参 考 文 献"
        incr_next = 1
        incr_kw = "附录A"
        offset_start = DM.get_anchor_position(ANCHOR)
        offset_end = super().render_template(ANCHOR, incr_next, incr_kw) - incr_next+1
        _style = DM.get_doc().styles['参考文献正文']
        for i in range(offset_start, offset_end):
            _p = DM.get_doc().paragraphs[i]
            _p.style = _style
            _p.paragraph_format.first_line_indent = Cm(-0.82)
        return offset_end
