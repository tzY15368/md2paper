from __future__ import annotations
import logging
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Cm
from typing import List, Union, Tuple
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_TAB_ALIGNMENT, WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from PIL import Image as PILImage
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree
import latex2mathml.converter

from md2paper.v2.backend.util import *


class Run():
    Normal = 1
    Italics = 2
    Bold = 4
    Formula = 8
    Superscript = 16
    Subscript = 32
    Reference = 64 | 1  # all `[ref]` as Normal style for now

    def __init__(self, text: str, style: int = 0, tabstop: bool = False, transform_required: bool = True) -> None:
        self.text = text
        self.bold = style & self.Bold != 0
        self.italics = style & self.Italics != 0
        self.formula = style & self.Formula != 0
        # transform rquired为true，则使用内置的xslt样式表，否则使用pandoc（需要pandoc依赖）
        self.__transform_required = transform_required
        self.subscript = style & self.Subscript != 0
        self.superscript = style & self.Superscript != 0
        self.__tabstop = tabstop

    def render_run(self, run):
        if self.formula and self.text:
            word_math = Formula.latex_to_word(
                self.text, transform_required=self.__transform_required)
            run._element.append(word_math)
        else:
            run.text = self.text
            run.bold = self.bold
            run.italic = self.italics
            run.font.subscript = self.subscript
            run.font.superscript = self.superscript

    @classmethod
    def get_tabstop(cls) -> Run:
        return Run("", tabstop=True)

    def is_tabstop(self) -> bool:
        return self.__tabstop


class BaseContent():
    # 在指定offset 填充paragraph，如果内部内容为多个paragraph，
    # 自动【向上】填充paragraph
    def render_paragraph(self, paragraph: Paragraph):
        raise NotImplementedError


class Text(BaseContent):
    # 换行会被该类内部自动处理

    two_chars_length = Cm(0.82)

    def __init__(self, raw_text: str = "", style: int = Run.Normal, force_style: str = '', first_line_indent=Cm(0.82)) -> None:
        self.runs: List[Run] = []
        self.force_style = force_style
        self.first_line_indent = first_line_indent
        if raw_text:
            self.runs.append(Run(raw_text, style))

    def add_run(self, run: Run) -> Text:
        self.runs.append(run)
        return self

    def add_hfill(self) -> Text:
        self.runs.append(Run.get_tabstop())
        return self

    def empty(self) -> bool:
        return len(self.runs) == 0

    def render_paragraph(self, position: Paragraph) -> int:
        if len(position.runs) != 0:
            logging.warning(
                "Text: existing content in paragraph", position.text)

        for run in self.runs:
            if not run.is_tabstop():
                run.render_run(position.add_run())
            else:
                # tabstop
                tab_stops = position.paragraph_format.tab_stops
                tab_stops.add_tab_stop(DM.margin_end, WD_TAB_ALIGNMENT.RIGHT)
        position.paragraph_format.first_line_indent = self.first_line_indent
        if self.force_style:
            position.style = DM.get_style(self.force_style)

    @classmethod
    def read(cls, txt: str) -> List[Text]:
        return [Text(i) for i in txt.split('\n')]


class ImageData():
    def __init__(self, src: str, alt: str, width_ratio=0) -> None:
        # 如果提供了0-1之间的width ratio，则会覆盖dpi设定，
        # 宽度1则图片宽约等于可编辑区域宽度，不等于纸张宽度。
        self.img_src = src
        self.img_alt = alt

        self.dpi = 360
        self.MAX_WIDTH_INCHES = 6

        if not self.img_src:
            self.size = (0, 0)
            self.size_inches = (0, 0)
            logging.debug("empty image, alt={}".format(self.img_alt))
            return

        img = PILImage.open(self.img_src)
        self.size = img.size
        img.close()

        img_size_ratio = self.size[0]/self.size[1]
        if width_ratio < 0 or width_ratio > 1:
            raise ValueError("invalid image width ratio, expecting range[0,1]")
        if width_ratio != 0:
            width_inches = self.MAX_WIDTH_INCHES*width_ratio
            height_inches = width_inches/img_size_ratio
            self.size_inches = (width_inches, height_inches)
        else:
            result = (self.size[0]/self.dpi, self.size[1]/self.dpi)
            if result[0] > self.MAX_WIDTH_INCHES:
                result = (self.MAX_WIDTH_INCHES,
                          self.MAX_WIDTH_INCHES/img_size_ratio)
            self.size_inches = result
        logging.debug(
            f"image size:{self.size_inches[0]},{self.size_inches[1]}")

    # returns width,height in Inches
    def get_size_in_doc(self) -> Tuple[Inches]:
        return map(Inches, self.size_inches)


class Image(BaseContent):
    # 子图由用户自行管理（大小、图题），这里默认只有一张图

    image_alt_style = "图名中文"

    def __init__(self, title: str, src: str) -> None:
        super().__init__()
        self.title = title
        self.src = src

    def set_image_data(self, data: ImageData) -> None:
        self.__image = data

    # 图片上下换行问题均由block管理，Image的render只负责图片和图题
    def render_paragraph(self, paragraph: Paragraph):
        p_text = paragraph

        p_text.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p_text.style = DM.get_style(self.image_alt_style)
        p_text.text = self.__image.img_alt

        if self.__image.img_src:
            p_img = paragraph.insert_paragraph_before()
            r = p_img.add_run()
            r.add_picture(self.__image.img_src, *
                          self.__image.get_size_in_doc())
            p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p_img.style = DM.get_style(self.image_alt_style)


class Formula(BaseContent):
    formula_alt_style = "图名中文"

    def __init__(self, title: str, formula: str, transform_required: bool = True) -> None:
        super().__init__()
        self.__title: str = title
        self.__formula: str = formula
        self.__transform_required = transform_required

    @classmethod
    def latex_to_word(cls, latex_input, transform_required=True):
        if not transform_required:
            return etree.fromstring(latex_input)
        mathml = latex2mathml.converter.convert(latex_input)
        tree = etree.fromstring(mathml)
        xslt = etree.parse(
            os.path.join(SRC_ROOT, 'backend', 'mml2omml.xsl')
        )
        transform = etree.XSLT(xslt)
        new_dom = transform(tree)
        return new_dom.getroot()

    def render_paragraph(self, paragraph: Paragraph):

        p_f = paragraph.insert_paragraph_before()
        table = DM.get_doc().add_table(rows=1, cols=3)
        p_f._p.addnext(table._tbl)
        DM.delete_paragraph(p_f)
        DM.delete_paragraph(paragraph)

        if self.__formula:
            cell_formula = table.rows[0].cells[1]
            cell_formula.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            Table.set_cell_border(cell_formula,
                                  top=Table.empty_border_tyle,
                                  bottom=Table.empty_border_tyle,
                                  start=Table.empty_border_tyle,
                                  end=Table.empty_border_tyle)
            par_formula = cell_formula.paragraphs[0]
            par_formula.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            r = par_formula.add_run()
            Run(self.__formula, Run.Formula,
                transform_required=self.__transform_required).render_run(r)

        # 标号
        cell_idx = table.rows[0].cells[2]
        cell_idx.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell_idx.paragraphs[0]
        p.text = self.__title
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        p.style = DM.get_style(self.formula_alt_style)

# row of table


class ListItem:
    def __init__(self, content_list: List[BaseContent]) -> None:
        self.content_list = content_list

    def render_paragraph(self, paragraph: Paragraph):
        for i, content in enumerate(self.content_list):
            if i == len(self.content_list)-1:
                p = paragraph
            else:
                p = paragraph.insert_paragraph_before()

            li_run = Run()
            if isinstance(content, Text):
                content.runs.insert(0, li_run)
            else:
                p2 = p.insert_paragraph_before()
                Text().add_run(li_run).render_paragraph(p2)
            content.render_paragraph(p)


class OrderedList(BaseContent):
    def __init__(self, item_list: List[ListItem]) -> None:
        self.item_list = item_list

    def render_paragraph(self, paragraph: Paragraph):
        for item in self.item_list:
            item.render_paragraph(paragraph)


class Row():
    def __init__(self, data: List[Text, str], top_border: bool = False) -> None:
        self.row: List[Text] = data
        self.has_top_border = top_border


class Table(BaseContent):
    black = '#000000'
    white = '#ffffff'
    empty_border_tyle = {'color': white}
    alt_style = "图名中文"

    def __init__(self, title: str, table: List[Row]) -> None:
        super().__init__()
        self.__auto_fit = True
        self.__columns_width: List[float] = []
        self.__title = title
        self.__table: List[Row] = table
        if len(table) < 1:
            raise ValueError("invalid table content")
        self.__cols = len(self.__table[0].row)
        self.__rows = len(self.__table)

    def set_columns_width(self, widths: List[float]):
        if len(widths) != self.__cols:
            raise ValueError(
                "invalid column width params, got {}, want {}", len(widths), self.__cols)
        self.__auto_fit = False
        self.__columns_width = widths

    def render_paragraph(self, paragraph: Paragraph):
        p = paragraph
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.style = DM.get_style(self.alt_style)
        p.text = self.__title
        table = DM.get_doc().add_table(rows=self.__rows, cols=self.__cols, style='Table Grid')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = table.allow_autofit = self.__auto_fit
        if not self.__auto_fit:
            for i in range(len(table.columns)):
                table.columns[i].width = Inches(
                    self.__columns_width[i] * DM.max_page_width_inches)
        paragraph._p.addnext(table._tbl)

        # 填充内容, 编辑表格样式
        for i, row in enumerate(self.__table):
            for j, cell_content in enumerate(row.row):
                cell = table.rows[i].cells[j]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if not self.__auto_fit:
                    cell.width = Inches(self.__columns_width[j] * 6)
                Table.set_cell_border(
                    cell,
                    top={
                        "val": 'single', 'color': self.white if not row.has_top_border else self.black},
                    bottom={"val": 'single', "color": self.white if i !=
                            self.__rows-1 else self.black},
                    start=Table.empty_border_tyle,
                    end=Table.empty_border_tyle
                )
                if cell_content == None:
                    if i == 0:
                        raise ValueError("invalid empty field in row 0")
                    else:
                        # 上一行同一列的cell
                        other_cell = table.rows[i-1].cells[j]
                        cell.merge(other_cell)
                        other_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                else:
                    p = cell.paragraphs[0]
                    cell_content.render_paragraph(p)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # https://stackoverflow.com/questions/33069697/how-to-setup-cell-borders-with-python-docx

    @classmethod
    def set_cell_border(cls, cell, **kwargs):
        """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # check for tag existnace, if none found, then create one
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        # list over all available tags
        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)

                # check for tag existnace, if none found, then create one
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)

                # looks like order of attributes is important
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)),
                                    str(edge_data[key]))


class Block():  # content
    # 每个block是多个image，formula，text的组合，内部有序
    Heading_1 = 1
    Heading_2 = 2
    Heading_3 = 3
    Heading_4 = 4

    def __init__(self) -> None:
        self.title: str = None
        self.level = -1
        self.__title_centered = False
        self.content_list: List[Union[Text, Image, Table, Formula]] = []
        self.sub_blocks: List[Block] = []

    # 由level决定标题的样式（heading1，2，3）
    def set_title(self, title: str, level: int, centered=False) -> Block:
        self.title = title
        if level not in range(0, 5):
            raise ValueError("invalid heading level")
        self.level = level
        self.__title_centered = centered
        return self

    # returns self
    def add_sub_block(self, block: Block) -> Block:
        self.sub_blocks.append(block)
        return self

    def get_last_sub_block(self) -> Block:
        if len(self.sub_blocks) == 0:
            raise ValueError("no available sub-blocks")
        return self.sub_blocks[-1]

    def add_content(self, *args: BaseContent):
        for content in args:
            if not isinstance(content, BaseContent):
                raise TypeError("expected BaseContent, got", type(content))
        self.content_list += args

    def render_template(self, paragraph: Paragraph = None):
        if not paragraph:
            # 最后开始append
            paragraph = DM.get_doc().add_paragraph()

        #
        if self.title:
            logging.debug(f"block(level={self.level}) title: {self.title}")
            p_title = paragraph.insert_paragraph_before()
            p_title.style = DM.get_style('Heading '+str(self.level))

            if self.__title_centered:
                p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # 如果是一级，给头上（标题前面）增加分页符
            if self.level == self.Heading_1:
                run = p_title.add_run()
                run.add_break(WD_BREAK.PAGE)
            title_run = p_title.add_run()
            title_run.text = self.title

        for i, content in enumerate(self.content_list):
            par = paragraph.insert_paragraph_before()
            content.render_paragraph(par)
            if i != len(self.content_list)-1:
                no_spacing_content = [Text, OrderedList]
                if type(self.content_list[i+1]) not in no_spacing_content:
                    paragraph.insert_paragraph_before().text = "Bbb"
                else:
                    if type(content) not in no_spacing_content:
                        paragraph.insert_paragraph_before().text = "aaa"

        for block in self.sub_blocks:
            par = paragraph.insert_paragraph_before()
            block.render_template(par)
