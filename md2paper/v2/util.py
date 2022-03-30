from typing import Union
import docx
from docx.shared import Inches
from docx.text.paragraph import Paragraph
import os
import logging
import lxml
import sys

sys.path.append('../..')
from io import BytesIO, StringIO 

from md2paper.md2paper import SRC_ROOT

class DocNotSetException(Exception):
    pass


class DocManager():
    __doc_target = None
    margin_end: Inches = None
    max_page_width_inches = 6
    @classmethod
    def get_style(cls,style_name:str):
        return cls.__doc_target.styles[style_name]

    @classmethod
    # doc_target: path-like string, file-like object or docx.Document
    def set_doc(cls, doc_target: Union[docx.Document, str, BytesIO]):
        if type(doc_target) == str:
            actual_path = os.path.join(SRC_ROOT, doc_target)
            logging.info(f"reading from template:{actual_path}")
            cls.__doc_target = docx.Document(actual_path)
        elif type(doc_target) == docx.Document:
            cls.__doc_target = doc_target
        elif type(doc_target) == BytesIO:
            cls.__doc_target = docx.Document(doc_target)
        else:
            raise TypeError(f"invalid doc target: expecting str or docx.Document type,\
                 got {type(doc_target)}")
        cls.__clear_tables()
        sec = cls.__doc_target.sections[0]
        cls.margin_end = docx.shared.Inches(
            sec.page_width.inches - (sec.left_margin.inches + sec.right_margin.inches))

    @classmethod
    def get_doc(cls) -> docx.Document:
        if not cls.__doc_target:
            raise DocNotSetException("doc target is not set, call DM.set_doc")
        return cls.__doc_target

    @classmethod
    def __clear_tables(cls):
        # delete all tables on startup as we don't need them
        for i in range(len(cls.get_doc().tables)):
            t = cls.get_doc().tables[0]._element
            t.getparent().remove(t)
            t._t = t._element = None

    @classmethod
    def delete_paragraph_by_index(cls, index):
        logging.debug(
            f"deleting idx={index} text={cls.get_paragraph(index).text}")
        p = cls.get_doc().paragraphs[index]._element
        p.getparent().remove(p)
        p._p = p._element = None

    @classmethod
    def delete_paragraph(cls, paragraph:Paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None
        

    @classmethod
    def get_anchor_position(cls, anchor_text: str, anchor_style_name="") -> int:
        # FIXME: 需要优化
        # 目前被设计成无状态的，只依赖template.docx文件以便测试，增加了性能开销
        # USE-WITH-CARE
        # 只靠标题的anchor-text找paragraph很容易找错，用的时候注意
        i = -1
        for _i, paragraph in enumerate(cls.get_doc().paragraphs):
            if anchor_text in paragraph.text:
                if (not anchor_style_name) or (paragraph.style.name == anchor_style_name):
                    i = _i
                    break

        if i == -1:
            raise ValueError(f"anchor `{anchor_text}` not found")
        return i + 1

    @classmethod
    def get_paragraph(cls, offset: int):
        return cls.get_doc().paragraphs[offset]

    # https://stackoverflow.com/questions/51360649/how-to-update-table-of-contents-in-docx-file-with-python-on-linux?rq=1
    @classmethod
    def update_toc(cls):
        namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        # add child to doc.settings element
        element_updatefields = lxml.etree.SubElement(
            cls.get_doc().settings.element, f"{namespace}updateFields"
        )
        element_updatefields.set(f"{namespace}val", "true")

    @classmethod
    def save(cls, out: Union[str, StringIO]):
        cls.__doc_target.save(out)


DM = DocManager
