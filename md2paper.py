from __future__ import annotations
from typing import Dict, Union,List
import docx
from docx.shared import Inches,Cm
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import lxml
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree
import latex2mathml.converter
import datetime

def latex_to_word(latex_input):
    mathml = latex2mathml.converter.convert(latex_input)
    tree = etree.fromstring(mathml)
    xslt = etree.parse(
        'mml2omml.xsl'
        )
    transform = etree.XSLT(xslt)
    new_dom = transform(tree)
    return new_dom.getroot()

class DocNotSetException(Exception):
    pass
class DocManager():
    __doc_target = None

    @classmethod
    def set_doc(cls,doc_target:docx.Document):
        cls.__doc_target = doc_target
        cls.__clear_tables()

    @classmethod
    def get_doc(cls)->docx.Document:
        if not cls.__doc_target:
            raise DocNotSetException("doc target is not set, call DM.set_doc")
        return cls.__doc_target

    @classmethod
    def __clear_tables(cls):
        # delete all tables on startup as we don't need them
        for i in range(len(cls.get_doc().tables)):
            t = cls.get_doc().tables[0]._element
            t.getparent().remove(t)
            t._t = t._element = None

    @classmethod
    def delete_paragraph_by_index(cls, index):
        p = cls.get_doc().paragraphs[index]._element
        p.getparent().remove(p)
        p._p = p._element = None    

    @classmethod
    def get_anchor_position(cls,anchor_text:str,anchor_style_name="")->int:
        # FIXME: 需要优化
        # 目前被设计成无状态的，只依赖template.docx文件以便测试，增加了性能开销
        # USE-WITH-CARE
        # 只靠标题的anchor-text找paragraph很容易找错，用的时候注意
        i = -1
        for _i,paragraph in enumerate(cls.get_doc().paragraphs):
            if anchor_text in paragraph.text:
                if (not anchor_style_name) or (paragraph.style.name == anchor_style_name):
                    i = _i
                    break
                        
        if i==-1: raise ValueError(f"anchor `{anchor_text}` not found") 
        return i + 1

    # https://stackoverflow.com/questions/51360649/how-to-update-table-of-contents-in-docx-file-with-python-on-linux?rq=1
    @classmethod
    def update_toc(cls):
        namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        # add child to doc.settings element
        element_updatefields = lxml.etree.SubElement(
            cls.get_doc().settings.element, f"{namespace}updateFields"
        )
        element_updatefields.set(f"{namespace}val", "true")

DM = DocManager

class BaseContent():

    # 在指定offset 【向上】填充paragraph，返回填充后最后一段的offset+1
    def render_paragraph(offset:int)->int:
        raise NotImplementedError
class Component():
    def __init__(self) -> None:
        self.__internal_text = Block()
    
    def get_internal_text(self)->Block:
        return self.__internal_text

    def get_default_location(self)->Block:
        return self.get_internal_text()

    def add_text(self, text:Union[str,List[BaseContent]], location:Block=None)->Union[BaseContent,None]:
        if not location:
            location = self.get_default_location()

        if type(text)==str:
            txt = Text.read(text)
            location.add_content(content_list=txt)
            if len(txt) > 0:
                return txt[-1]
            else:
                return None

        elif type(text)==list and sum([1 if isinstance(i,BaseContent) else 0 for i in text])==len(text):
            location.add_content(content_list=text)
            if len(text) > 0:
                return text[-1]
            else:
                return None
        else:
            raise TypeError("expected text/List[BaseContent]")

    # anchor_text: 用于找到插入段落位置
    # incr_next: 用于在插入新内容后往后删老模板当前段内容，
    # 直到删除到incr_kw往前incr_next个paragraph
    # incr_kw：见上面incr_next
    def render_template(self, anchor_text:str,  incr_next:int, incr_kw, anchor_style_name="")->int:
        offset = DM.get_anchor_position(anchor_text=anchor_text,anchor_style_name=anchor_style_name)
        while not incr_kw in DM.get_doc().paragraphs[offset+incr_next].text\
                 and (offset+incr_next)!=(len(DM.get_doc().paragraphs)-1):
            DM.delete_paragraph_by_index(offset)
            #print('deleted 1 line for anchor',anchor_text)
        return self.__internal_text.render_template(offset)


class Run():
    Normal = 1
    Italics = 2
    Bold = 4
    Formula = 8
    Superscript = 16
    Subscript = 32
    def __init__(self,text:str,style:int=0,tabstop:bool=False) -> None:
        self.text = text
        self.bold = style & self.Bold != 0
        self.italics = style & self.Italics != 0
        self.formula = style & self.Formula != 0
        self.subscript = style & self.Subscript != 0
        self.superscript = style & self.Superscript != 0
        self.__tabstop = tabstop
    
    def render_run(self,run):
        if not self.formula:
            run.text = self.text
            run.bold = self.bold
            run.italic = self.italics
            run.font.subscript = self.subscript
            run.font.superscript = self.superscript

        else:
            word_math = latex_to_word(self.text)
            run._element.append(word_math)

    @classmethod
    def get_tabstop(cls)->Run:
        return Run("",tabstop=True)

    def is_tabstop(self)->bool:
        return self.__tabstop
class Text(BaseContent):
    # 换行会被该类内部自动处理

    def __init__(self, raw_text:str="",style:int=Run.Normal) -> None:
        self.__runs:List[Run] = []
        if raw_text:
            self.__runs.append(Run(raw_text,style))

    def add_run(self, run:Run)->Text:
        self.__runs.append(run)
        return self

    def add_hfill(self)->Text:
        self.__runs.append(Run.get_tabstop())
        return self

    def render_paragraph(self, offset: int) -> int:
        new_offset = offset
        p = DM.get_doc().paragraphs[new_offset].insert_paragraph_before()
        for run in self.__runs:
            if not run.is_tabstop():
                run.render_run(p.add_run())
            else:
                # https://stackoverflow.com/questions/58656450/how-to-use-tabletop-by-python-docx
                sec = DM.get_doc().sections[0]
                margin_end = docx.shared.Inches(
                    sec.page_width.inches - (sec.left_margin.inches + sec.right_margin.inches))
                tab_stops = p.paragraph_format.tab_stops
                # adding new tab stop, to the end point, and making sure that it's `RIGHT` aligned.
                tab_stops.add_tab_stop(margin_end, docx.enum.text.WD_TAB_ALIGNMENT.RIGHT)
                
        p.paragraph_format.first_line_indent = Cm(0.82)
        new_offset = new_offset + 1
        return new_offset

    @classmethod
    def read(cls, txt:str)->List[Text]:
        return [Text(i) for i in txt.split('\n')]

class ImageData():
    def __init__(self,src:str,alt:str) -> None:
        self.img_src = src
        self.img_alt = alt

class Image(BaseContent):
    def __init__(self,data:List[ImageData]) -> None:
        super().__init__()
        self.__images = data

    def render_paragraph(self, offset: int) -> int:
        new_offset = offset
        for img in self.__images:
            p = DM.get_doc().paragraphs[new_offset].insert_paragraph_before()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER 
            p.style = DM.get_doc().styles['图名中文']
            # 先换一行
            r0 = p.add_run()
            r0.add_break(WD_BREAK.LINE)
            r = p.add_run()
            r.add_picture(img.img_src,width=Inches(4.0), height=Inches(4.7))
            r2 = p.add_run()
            r2.add_break(WD_BREAK.LINE)
            r3 = p.add_run()
            r3.add_text("\n"+img.img_alt)
            # 结尾再换
            r4 = p.add_run()
            r4.add_break(WD_BREAK.LINE)

            new_offset = new_offset + 1
        
        return new_offset

class Row():
    def __init__(self,data:List[str],top_border:bool=False) -> None:
        self.row:List[str] = data
        self.has_top_border = top_border

class Formula(BaseContent):
    def __init__(self,title:str,formula:str) -> None:
        super().__init__()
        self.__title:str = title
        self.__formula:str = formula
        
    def render_paragraph(self,offset: int) -> int:
        new_offset = offset
        p = DM.get_doc().paragraphs[new_offset].insert_paragraph_before()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table = DM.get_doc().add_table(rows=1,cols=3)
        p._p.addnext(table._tbl)
        DM.delete_paragraph_by_index(new_offset)

        # 公式cell
        cell_formula = table.rows[0].cells[1]
        cell_formula.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        empty = {'color':Table.white}
        Table.set_cell_border(
            cell_formula,
            top=empty,
            bottom=empty,
            start=empty,
            end=empty
        )
        _p = cell_formula.paragraphs[0]
        _p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = _p.add_run()
        Run(self.__formula,Run.Formula).render_run(r)

        # 标号cell
        cell_idx = table.rows[0].cells[2]
        cell_idx.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_idx_p = cell_idx.paragraphs[0]
        cell_idx_p.text = self.__title
        cell_idx_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        return new_offset


class Table(BaseContent):
    # table的最后一行默认有下边框，剩下依靠row的top-border自行决定
    black = "#000000"
    white = "#ffffff"
    def __init__(self,title:str, table:List[Row]) -> None:
        super().__init__()
        self.__title = title
        self.__table:List[Row] = table
        if len(table) < 1:
            raise ValueError("invalid table content")
        self.__cols = len(self.__table[0].row)
        self.__rows = len(self.__table)
    
    def render_paragraph(self, offset: int) -> int:
        new_offset = offset
        p1 = DM.get_doc().paragraphs[new_offset].insert_paragraph_before()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER 
        p1.style = DM.get_doc().styles['图名中文']
        # 先换一行
        r0 = p1.add_run()
        r0.add_break(WD_BREAK.LINE)
        
        r1 = p1.add_run()
        r1.add_text(self.__title)
        new_offset = new_offset + 1

        p2 = DM.get_doc().paragraphs[new_offset].insert_paragraph_before()
        # 表格自带一个换行？
        table = DM.get_doc().add_table(rows = self.__rows, cols = self.__cols, style='Table Grid')
        # 将table挪到paragrpah里
        p2._p.addnext(table._tbl)
        # 挪完删掉paragraph
        DM.delete_paragraph_by_index(new_offset)
        
        new_offset = new_offset + 1
        
        # 结尾再换
        p1 = DM.get_doc().paragraphs[new_offset].insert_paragraph_before()
        new_offset = new_offset + 1

        # 填充内容, 编辑表格样式
        for i,row in enumerate(self.__table):
            for j,cell_str in enumerate(row.row):
                cell = table.rows[i].cells[j]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                Table.set_cell_border(
                    cell,
                    top={"val":'single','color':self.white if not row.has_top_border else self.black},
                    bottom = {"val":'single', "color":self.white if i!=self.__rows-1 else self.black},
                    start={"color":self.white},
                    end={"color":self.white}
                )
                if cell_str == None:
                    if i == 0:
                        raise ValueError("invalid empty field in row 0")
                    else:
                        # 上一行同一列的cell
                        other_cell = table.rows[i-1].cells[j]
                        cell.merge(other_cell)
                        other_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                else:
                    p = cell.paragraphs[0]
                    p.text = cell_str
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.style = DM.get_doc().styles['图名中文']

        return new_offset

    # https://stackoverflow.com/questions/33069697/how-to-setup-cell-borders-with-python-docx
    @classmethod
    def set_cell_border(cls, cell, **kwargs):
        """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # check for tag existnace, if none found, then create one
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        # list over all available tags
        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)

                # check for tag existnace, if none found, then create one
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)

                # looks like order of attributes is important
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))

class Metadata(Component):
    school: str = None
    major: str = None
    name: str = None
    number: str = None
    teacher: str = None
    auditor: str = None
    __finish_date: str = None
    title_zh_CN: str = None
    title_en: str = None

    @property
    def finish_date(self):
        if self.__finish_date:
            return self.__finish_date
        else:
            now = datetime.datetime.now()
            return now.strftime("%Y年%m月%d日")

    @finish_date.setter
    def finish_date(self, value):
        self.__finish_date = value

    def __fill_blank(self, blank_length:int, data:str)->str:
        """
        填充诸如 "学 生 姓 名：______________"的域
        **一个中文算两个字符
        **中文和数字之间会多一个空格
        """
        def get_data_len(data:str)->int:
            # 判断是否中文, 一个中文算两个字符
            def is_zh_CN(char)->bool:
                return '\u4e00' <= char <= '\u9fa5'
            len = 0
            prev_char_type = False
            prev_char = None
            for char in data:
                current_char_type = is_zh_CN(char)
                len += 1
                if current_char_type:
                    len += 1
                if current_char_type != prev_char_type and prev_char!=None:
                    len += 0.5
                prev_char_type = current_char_type
                prev_char = char
            return int(len)

        head_length = int((blank_length - get_data_len(data)) /2)
        if head_length <0:
            raise ValueError("值过长")
        content = " " * head_length + data + " " * (blank_length-get_data_len(data)-head_length)
        #print(data,get_data_len(data))
        return content

    def render_template(self):
        # 目前只支持论文，不支持翻译！！

        title_mapping = {
            'zh_CN': 4,
            'en': 5,
        }
        DM.get_doc().paragraphs[title_mapping['zh_CN']].runs[0].text = self.title_zh_CN
        DM.get_doc().paragraphs[title_mapping['en']].runs[0].text = self.title_en

        line_mapping = {
            15: self.school,
            16: self.major,
            17: self.name,
            18: self.number,
            19: self.teacher,
            20: self.auditor,
            21: self.finish_date
        }
        BLANK_LENGTH = 23
        for line_no in line_mapping:
            if line_mapping[line_no] == None:
                continue
            #print(len(DM.get_doc().paragraphs[line_no].runs[-1].text))
            DM.get_doc().paragraphs[line_no].runs[-1].text = self.__fill_blank(BLANK_LENGTH,line_mapping[line_no])

class Abstract(Component):

    def __init__(self) -> None:
        self.__text_en:Block = Block()
        self.__text_zh_CN:Block = Block()
        self.__keyword_en:str = None
        self.__keyword_zh_CN:str = None
        self.__title_zh_CN = ""
        self.__title_en = ""

    # 中文title被忽略，因为模板中没用上
    def set_title(self,zh_CN:str, en:str):
        self.__title_zh_CN = zh_CN
        self.__title_en = en

    def set_keyword(self, zh_CN:List[str],en:List[str]):
        SEPARATOR = "；"
        self.__keyword_en = SEPARATOR.join(en)
        self.__keyword_zh_CN = SEPARATOR.join(zh_CN)

    def add_text(self, zh_CN:str,en:str):
        self.__text_en.add_content(content_list=Text.read(en))
        self.__text_zh_CN.add_content(content_list=Text.read(zh_CN))

    def render_template(self)->int:
        # 64开始是摘要正文
        abs_cn_start = 64
        abs_cn_end = self.__text_zh_CN.render_block(abs_cn_start)
        
        while not DM.get_doc().paragraphs[abs_cn_end+2].text.startswith("关键词："):
            DM.delete_paragraph_by_index(abs_cn_end+1)
        
        # cn kw
        kw_cn_start = abs_cn_end + 2
        DM.get_doc().paragraphs[kw_cn_start].runs[1].text = self.__keyword_zh_CN
        
        # en start

        en_title_start = kw_cn_start+4
        DM.get_doc().paragraphs[en_title_start].runs[1].text = self.__title_en

        en_abs_start = en_title_start + 3
        en_abs_end = self.__text_en.render_block(en_abs_start)-1

        # https://stackoverflow.com/questions/61335992/how-can-i-use-python-to-delete-certain-paragraphs-in-docx-document
        while not DM.get_doc().paragraphs[en_abs_end+2].text.startswith("Key Words："):
            DM.delete_paragraph_by_index(en_abs_end+1)

        # en kw
        kw_en_start = en_abs_end +2

        # https://github.com/python-openxml/python-docx/issues/740
        delete_num = len(DM.get_doc().paragraphs[kw_en_start].runs) - 4
        for run in reversed(list(DM.get_doc().paragraphs[kw_en_start].runs)):
            DM.get_doc().paragraphs[kw_en_start]._p.remove(run._r)
            delete_num -= 1
            if delete_num < 1:
                break
        
        
        DM.get_doc().paragraphs[kw_en_start].runs[3].text = self.__keyword_en
        return kw_en_start+1
class Conclusion(Component):
    def render_template(self) -> int:
        ANCHOR = "结    论（设计类为设计总结）"
        incr_next = 3
        incr_kw = "参 考 文 献"
        return super().render_template(ANCHOR, incr_next, incr_kw)

class MainContent(Component): # 正文

    def __init__(self) -> None:
        super().__init__()
        self.__prev:Dict[str,Block] = {
            'blk':None,
            'chapter':None,
            'section':None,
            'subsection':None
        }
        
    def get_last_subblock(self)->Block:
        if 'blk' not in self.__prev:
            raise ValueError("last blk is not yet set")
        return self.__prev['blk']

    # get_last_subblock
    def get_default_location(self) -> Block:
        if 'blk' not in self.__prev:
            raise ValueError("last blk is not yet set")
        return self.__prev['blk']

    # add_chapter returns the added chapter
    def add_chapter(self,title:str)->Block:
        new_chapter = Block()

        # set prev pointers
        self.__prev['blk'] = new_chapter
        self.__prev['chapter'] = new_chapter

        new_chapter.set_title(title,Block.heading_1)
        return self.get_internal_text().add_sub_block(new_chapter)
    
    # add_section returns the added section
    def add_section(self, title:str, chapter:Block=None)->Block:
        if not chapter:
            if self.__prev.get('chapter'):
                chapter = self.__prev['chapter']
            else:
                raise KeyError("chapter must be initialized before section")

        new_section = Block()

        # set prev pointers
        self.__prev['blk'] = new_section
        self.__prev['section'] = new_section

        new_section.set_title(title,Block.heading_2)
        return chapter.add_sub_block(new_section)

    # add_subsection returns the added subsection
    def add_subsection(self, title:str,section:Block=None)->Block:
        if not section:
            if 'section' in self.__prev:
                section = self.__prev['section']
            else:
                raise KeyError("section must be initialized before subsection")
        
        new_subsection = Block()
        
        # set prev pointers
        self.__prev['blk'] = new_subsection
        self.__prev['subsection'] = new_subsection

        new_subsection.set_title(title,Block.heading_3)
        return section.add_sub_block(new_subsection)

    # 由于无法定位正文，需要先生成引言，再用引言返回的offset
    def render_template(self) -> int:
        anchor_text = "1  正文格式说明"
        anchor_style = "Heading 1"
        incr_next = 3
        incr_kw = "结    论（设计类为设计总结"
        #此处没有覆盖原有内容，因此还需要删去原有的大标题 1 正文格式……
        offset = super().render_template(anchor_text,incr_next,incr_kw,anchor_style_name=anchor_style)
        
        line_delete_count = 1
        pos = DM.get_anchor_position(anchor_text,anchor_style_name=anchor_style) - 1
        for i in range(line_delete_count):
            DM.delete_paragraph_by_index(pos)
        return offset - line_delete_count

class Introduction(Component): #引言 由于正文定位依赖引言，如果没写引言，依旧会生成引言，最后删掉
    def render_template(self) -> int:
        anchor_text = "引    言"
        incr_next = 2
        incr_kw = "正文格式说明"
        anchor_style_name = "Heading 1"
        return super().render_template(anchor_text, incr_next, incr_kw, anchor_style_name=anchor_style_name)


class References(Component): #参考文献
    def render_template(self) -> int:
        ANCHOR = "参 考 文 献"
        incr_next = 1
        incr_kw = "附录A"
        offset_start = DM.get_anchor_position(ANCHOR)
        offset_end = super().render_template(ANCHOR, incr_next, incr_kw) -incr_next+1
        _style = DM.get_doc().styles['参考文献正文']
        for i in range(offset_start,offset_end):
            DM.get_doc().paragraphs[i].style = _style
        return offset_end
        

class Appendixes(Component): #附录abcdefg, 是一种特殊的正文
    def __init__(self) -> None:
        super().__init__()

    def get_default_location(self)->Block:
        return self.get_internal_text().get_last_sub_block()

    def add_appendix(self, appendix_title:str)->Block:
        new_appendix = Block()
        new_appendix.set_title(appendix_title,level=Block.heading_1)
        self.get_internal_text().add_sub_block(new_appendix)
        return new_appendix

    # returns the added appendix
    def add_appendix2(self, appendix_title:str,appendix_content:str)->Block:
        new_appendix = Block()
        new_appendix.set_title(appendix_title,level=Block.heading_1)
        new_appendix.add_content(content_list=Text.read(appendix_content))
        self.get_internal_text().add_sub_block(new_appendix)

    def render_template(self) -> int:
        anchor_text = "附录A"
        anchor_style_name = "Heading 1"
        incr_next = 1
        incr_kw = "修改记录"
        offset = super().render_template(anchor_text, incr_next, incr_kw, anchor_style_name)
        #此处没有覆盖原有内容，因此还需要删去原有的附录a那一页的3段

        line_delete_count = 1
        pos = DM.get_anchor_position(anchor_text=anchor_text)-1
        for i in range(line_delete_count):
            DM.delete_paragraph_by_index(pos)
        return offset - line_delete_count
    

class ChangeRecord(Component): #修改记录
    def render_template(self) -> int:
        ANCHOR = "修改记录"
        ANCHOR_STYLE = "Heading 1"
        incr_next = 0
        incr_kw = "致    谢"
        return super().render_template(ANCHOR,incr_next,incr_kw,anchor_style_name=ANCHOR_STYLE)

class Acknowledgments(Component): #致谢
    def render_template(self) -> int:
        ANCHOR = "致    谢"
        incr_next = 0

        #hack: 致谢已经到论文末尾，因此用无法匹配上的字符串直接让他删到最后一行
        incr_kw = "/\,.;'" 
        return super().render_template(ANCHOR,incr_next,incr_kw)
    

"""
每一章是一个chapter，
每个chapter内标题后可以直接跟block或section，
每个section内标题后可以直接跟block或subsection，
每个subsection内标题后可以跟block
"""
class Block(): #content
    # 每个block是多个image，formula，text的组合，内部有序
    heading_1 = 1
    heading_2 = 2
    heading_3 = 3
    heading_4 = 4

    def __init__(self) -> None:
        self.__title:str = None
        self.__content_list:List[Union[Text,Image,Table,Formula]] = []
        self.__sub_blocks:List[Block] = []
        self.__id:int = None

    def set_id(self, id:int):
        self.__id = id

    # 由level决定标题的样式（heading1，2，3）
    def set_title(self,title:str,level:int) -> Block:
        self.__title = title
        if level not in range(0,5):
            raise ValueError("invalid heading level")
        self.__level = level
        return self
    
    def add_sub_block(self,block:Block)->Block:
        self.__sub_blocks.append(block)
        return block

    def get_last_sub_block(self)->Block:
        if len(self.__sub_blocks)==0:
            raise ValueError("no available sub-blocks")
        return self.__sub_blocks[-1]

    def add_content(self,content:Union[Text,Image,Table,Formula]=None,
            content_list:Union[List[Text],List[Image],List[Table],List[Formula]]=[]) -> Block:
        if content:
            self.__content_list.append(content)
        for i in content_list:
            self.__content_list.append(i)
        #print('added content with len',len(content_list),content_list[0].raw_text,id(self))
        return self

    # render_template是基于render_block的api，增加了嵌套blocks的渲染 以支持递归生成章节/段落，
    # 同时增加了对段落标题和段落号的支持
    # 顺序：先title，再自己的content-list，再自己的sub-block
    def render_template(self, offset:int)->int:
        new_offset = offset

        # 如果是一级，给头上（标题前面）增加分页符
        if self.__title and self.__level == self.heading_1:
            p = DM.get_doc().paragraphs[new_offset].insert_paragraph_before()
            run = p.add_run()
            run.add_break(WD_BREAK.PAGE)
            new_offset = new_offset + 1

        if self.__title:
            #print('title:',self.__title,'level:',self.__level)
            p_title = DM.get_doc().paragraphs[new_offset].insert_paragraph_before()
            p_title.style = DM.get_doc().styles['Heading '+str(self.__level)]
            p_title.add_run()
            title_idx = "" if not self.__id else str(self.__id) + "  "
            p_title.runs[0].text= title_idx + self.__title
            new_offset = new_offset + 1
        
        new_offset = self.render_block(new_offset)

        #print('has',len(self.__sub_blocks),"sub-blocks")
        for i,block in enumerate(self.__sub_blocks):
            new_offset = block.render_template(new_offset) 

        return new_offset


    # render_block是最底层的api，只将自己的content-list加到已有文档给定位置
    # render_block takes the desired paragraph position's offset,
    # renders the block with native elements: text, image and formulas,
    # and returns the final paragraph's offset
    
    def render_block(self, offset:int)->int:
        if not self.__content_list:
            return offset
        new_offset = offset
        for content in self.__content_list:
            new_offset = content.render_paragraph(new_offset)
        return new_offset

if __name__ == "__main__":
    
    doc = docx.Document("毕业设计（论文）模板-docx.docx")
    DM.set_doc(doc)
    meta = Metadata()
    meta.school = "电子信息与电气工程"
    meta.number = "201800000"
    meta.auditor = "张三"
    meta.finish_date = "1234年5月6号"
    meta.teacher = "里斯"
    meta.major = "计算机科学与技术"
    meta.title_en = "what is this world"
    meta.title_zh_CN = "这是个怎样的世界"
    meta.render_template()

    abs = Abstract()
    a = """CommonMark中并未定义普通文本高亮。
CSDN中支持通过==文本高亮==，实现文本高亮

如果你使用的markdown编辑器不支持该便捷设置，可通过HTML标签<mark>实现：
语法：<mark>文本高亮<mark>
效果：文本高亮"""
    b = ['您配吗','那匹马','进来看是否']
    
    c = """Any subsequent access to the "deleted" paragraph object will raise AttributeError, so you should be careful not to keep the reference hanging around, including as a member of a stored value of Document.paragraphs.
The reason it's not in the library yet is because the general case is much trickier, in particular needing to detect and handle the variety of linked items that can be present in a paragraph; things like a picture, a hyperlink, or chart etc.
But if you know for sure none of those are present, these few lines should get the job done."""
    
    d = ['abc','def','gh']

    abs.add_text(a,c)
    abs.set_keyword(b,d)
    abs.set_title(meta.title_zh_CN,meta.title_en)
    abs.render_template()

    intro = Introduction()
    t = """这样做违反了Liskov替代原则。换句话说，这是一个可怕的想法，B不应该是A的子类型。我只是感兴趣：您为什么要这样做？@delnan出于某种原因每当有人提到我总是想到Who Doctor的Blinovitch限制效应时。
现在就称其为好奇心。我感谢警告，但我仍然感到好奇。
一个用例是，如果您要使用Django库公开的Form类，但不包含其字段之一。在Django中，表单字段是由某些类属性定义的。例如，请参阅此SO问题。"""
    intro.add_text(t)
    txt = Text().add_run(Run("a"))
    txt.add_run(Run("bbb",Run.Italics|Run.Subscript))
    txt.add_run(Run(" and then A",Run.Bold))
    txt.add_run(Run("2",Run.Superscript))
    intro.add_text([txt])
    intro.render_template()

    mc = MainContent()
    c1 = mc.add_chapter("第一章 刘姥姥")
    s1 = mc.add_section("1.1 asdfasdf",chapter=c1)
    s2 = mc.add_section("1.2 bbbb",chapter=c1)
    h = """目前的娛樂型電腦螢幕市場，依照玩家的需求大致可以分為兩大勢力：一派是主打對戰類型
    的電競玩家、另一派則主打追劇的多媒體影音玩家。前者需要需要高更新率的螢幕，在分秒必爭的對戰中搶得先機；後者則需要較高的解析度以及HDR的顯示內容，好用來欣賞畫面的每一個細節。"""
    mc.add_text(h,location=c1)
    mc.add_text(h,location=s2)
    c2 = mc.add_chapter("第二章 菜花")
    s3 = mc.add_section("2.1 aaa",chapter=c2)
    ss1 = mc.add_subsection("2.1.1 asdf",section=s3)
    txt = mc.add_text(h,location=ss1)
    txt.add_run(Run("this should be bold",Run.Bold))
    txt.add_run(Run("italic and bold",Run.Italics|Run.Bold))
    images = Image([
        ImageData("classes.png","图1：these are the classes"),
        ImageData("classes.png","图2:asldkfja;sldkf")
    ])
    formula = Formula("公式3.4",r"\sum_{i=1}^{10}{\frac{\sigma_{zp,i}}{E_i} kN")
    more_text = Text().add_run(Run("only italics",Run.Italics))
    mc.add_text([images,formula,more_text])

    c3 = mc.add_chapter("第三章 大观园")
    mc.add_text(t,location=c3)
    data = [
        Row(['第一章','第二章','第三章'],top_border=True),
        Row(['刘姥姥初试钢铁侠','刘姥姥初试大不净者','刘姥姥倒拔绿巨人'],top_border=True),
        Row(['刘姥姥初试惊奇队长',None,'刘姥姥菜花染诸神']),
        Row(['菜花反噬！','天地乖离菜花之星','重启刘姥姥菜花宇宙'],top_border=True)
    ]
    table = Table("表1 刘姥姥背叛斯大林",data)
    mc.add_text([table, Text("wtf is this?")],location=c3)
    mc.render_template()

    conc = Conclusion()
    e = """如果代码中出现太多的条件判断语句的话，代码就会变得难以维护和阅读。 这里的解决方案是将每个状态抽取出来定义成一个类。
这里看上去有点奇怪，每个状态对象都只有静态方法，并没有存储任何的实例属性数据。 实际上，所有状态信息都只存储在 Connection 实例中。 
在基类中定义的 NotImplementedError 是为了确保子类实现了相应的方法。 这里你或许还想使用8.12小节讲解的抽象基类方式。
设计模式中有一种模式叫状态模式，这一小节算是一个初步入门！"""
    #conc.set_conclusion(e)
    conc.add_text(e)
    conc.add_text([Text().add_run(Run("additional italics text",Run.Italics)), Image([ImageData("classes.png","图3：这是classes")])])
    conc.render_template()

    ref = References()
    h = """[1] 国家标准局信息分类编码研究所.GB/T 2659-1986 世界各国和地区名称代码[S]//全国文献工作标准化技术委员会.文献工作国家标准汇编:3.北京:中国标准出版社,1988:59-92. 
[2] 韩吉人.论职工教育的特点[G]//中国职工教育研究会.职工教育研究论文集.北京:人民教育出版社,1985:90-99. """
    ref.add_text(h)
    ref.render_template()

    ack = Acknowledgments()
    f = """肾衰竭（Kidney failure）是一种终末期的肾脏疾病，此时肾脏的功能会低于其正常水平的15%。由于透析会严重影响患者的生活质量，肾移植一直是治疗肾衰竭的理想方式。但肾脏供体一直处于短缺状态，移植需等待时间约为5-10年。近日，据一篇发表于《美国移植杂志》的文章，阿拉巴马大学伯明翰分校的科学家首次成功将基因编辑猪的肾脏成功移植给一名脑死亡的人类接受者。
研究使用的供体猪的10个关键基因经过了基因编辑，其肾脏的功能更适合人体，且植入人体后引发的免疫排斥反应更轻微。研究人员首先对异种供体和接受者进行交叉配型测试。经配对后，研究人员将肾脏移植入脑死亡接受者的肾脏对应的解剖位置，与肾动脉、肾静脉和输尿管相连接。移植后，他们为患者进行了常规的免疫抑制治疗。目前，肾脏已在患者体内正常工作77小时。该研究按照1期临床试验标准进行，完全按人类供体器官的移植标准实施。研究显示，异种移植的发展在未来或可缓解世界范围器官供应压力。"""
    ack.add_text(f)
    ack.render_template()

    cha = ChangeRecord()
    g = """在无线传能技术中，非辐射无线传能（即电磁感应充电）可以高效传输能量，但有效传输距离被限制在收发器尺寸的几倍之内；而辐射无线传能（如无线电、激光）虽然可以远距离传输能量，但需要复杂的控制机制来跟踪移动的能量接收器。
近日，同济大学电子与信息工程学院的研究团队通过理论和实验证明，"""
    cha.add_text(g)
    cha.render_template()

    apd = Appendixes()
    apd.add_appendix("附录A")
    apd.add_text([Text("啊哈哈哈鸡汤来喽")])
    apd.add_appendix("附录B")
    apd.add_text([Text("直接来吧")])
    apd.render_template()

    #DM.update_toc()
    doc.save("out.docx")

    