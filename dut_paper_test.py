from dut_paper import *

import docx


if __name__ == "__main__":
    doc = docx.Document("毕业设计（论文）模板-docx.docx")
    DM.set_doc(doc)
    meta = Metadata()
    meta.school = "电子信息与电气工程"
    meta.number = "201800000"
    meta.auditor = "张三"
    meta.finish_date = "1234年5月6号"
    meta.teacher = "里斯"
    meta.major = "计算机科学与技术"
    meta.title_en = "what is this world"
    meta.title_zh_CN = "这是个怎样的世界"
    meta.render_template()

    abs = Abstract()
    a = """CommonMark中并未定义普通文本高亮。
CSDN中支持通过==文本高亮==，实现文本高亮

如果你使用的markdown编辑器不支持该便捷设置，可通过HTML标签<mark>实现：
语法：<mark>文本高亮<mark>
效果：文本高亮"""
    b = ['您配吗','那匹马','进来看是否']
    
    c = """Any subsequent access to the "deleted" paragraph object will raise AttributeError, so you should be careful not to keep the reference hanging around, including as a member of a stored value of Document.paragraphs.
The reason it's not in the library yet is because the general case is much trickier, in particular needing to detect and handle the variety of linked items that can be present in a paragraph; things like a picture, a hyperlink, or chart etc.
But if you know for sure none of those are present, these few lines should get the job done."""
    
    d = ['abc','def','gh']

    abs.add_text(a,c)
    abs.set_keyword(b,d)
    abs.set_title(meta.title_zh_CN,meta.title_en)
    abs.render_template()

    intro = Introduction()
    t = """这样做违反了Liskov替代原则。换句话说，这是一个可怕的想法，B不应该是A的子类型。我只是感兴趣：您为什么要这样做？@delnan出于某种原因每当有人提到我总是想到Who Doctor的Blinovitch限制效应时。
现在就称其为好奇心。我感谢警告，但我仍然感到好奇。
一个用例是，如果您要使用Django库公开的Form类，但不包含其字段之一。在Django中，表单字段是由某些类属性定义的。例如，请参阅此SO问题。"""
    intro.add_text(t)
    txt = Text().add_run(Run("a"))
    txt.add_run(Run("bbb",Run.Italics|Run.Subscript))
    txt.add_run(Run(" and then A",Run.Bold))
    txt.add_run(Run("2",Run.Superscript))
    intro.add_text([txt])
    intro.render_template()


    mc = MainContent()
    c1 = mc.add_chapter("第一章 刘姥姥")
    s1 = mc.add_section("1.1 asdfasdf",chapter=c1)
    s2 = mc.add_section("1.2 bbbb",chapter=c1)
    h = """目前的娛樂型電腦螢幕市場，依照玩家的需求大致可以分為兩大勢力：一派是主打對戰類型
    的電競玩家、另一派則主打追劇的多媒體影音玩家。前者需要需要高更新率的螢幕，在分秒必爭的對戰中搶得先機；後者則需要較高的解析度以及HDR的顯示內容，好用來欣賞畫面的每一個細節。"""
    mc.add_text(h,location=c1)
    mc.add_text(h,location=s2)
    c2 = mc.add_chapter("第二章 菜花")
    s3 = mc.add_section("2.1 aaa",chapter=c2)
    ss1 = mc.add_subsection("2.1.1 asdf",section=s3)
    txt = mc.add_text(h,location=ss1)
    txt.add_run(Run("this should be bold",Run.Bold))
    txt.add_run(Run("italic and bold",Run.Italics|Run.Bold))
    images = Image([
        ImageData("classes.png","图1：these are the classes"),
        ImageData("classes.png","图2:asldkfja;sldkf")
    ])
    formula = Formula("公式3.4",r"\sum_{i=1}^{10}{\frac{\sigma_{zp,i}}{E_i} kN")
    more_text = Text().add_run(Run("only italics",Run.Italics))
    mc.add_text([images,formula,more_text])

    c3 = mc.add_chapter("第三章 大观园")
    mc.add_text(t,location=c3)
    data = [
        Row(['第一章','第二章','第三章'],top_border=True),
        Row(['刘姥姥初试钢铁侠','刘姥姥初试大不净者','刘姥姥倒拔绿巨人'],top_border=True),
        Row(['刘姥姥初试惊奇队长',None,'刘姥姥菜花染诸神']),
        Row(['菜花反噬！','天地乖离菜花之星','重启刘姥姥菜花宇宙'],top_border=True)
    ]
    table = Table("表1 刘姥姥背叛斯大林",data)
    mc.add_text([table, Text("wtf is this?")],location=c3)
    mc.render_template()


    conc = Conclusion()
    e = """如果代码中出现太多的条件判断语句的话，代码就会变得难以维护和阅读。 这里的解决方案是将每个状态抽取出来定义成一个类。
这里看上去有点奇怪，每个状态对象都只有静态方法，并没有存储任何的实例属性数据。 实际上，所有状态信息都只存储在 Connection 实例中。 
在基类中定义的 NotImplementedError 是为了确保子类实现了相应的方法。 这里你或许还想使用8.12小节讲解的抽象基类方式。
设计模式中有一种模式叫状态模式，这一小节算是一个初步入门！"""
    conc.add_text(e)
    conc.add_text([Text().add_run(Run("additional italics text",Run.Italics)), Image([ImageData("classes.png","图3：这是classes")])])
    conc.render_template(override_title="设计总结")

    ref = References()
    h = """[1] 国家标准局信息分类编码研究所.GB/T 2659-1986 世界各国和地区名称代码[S]//全国文献工作标准化技术委员会.文献工作国家标准汇编:3.北京:中国标准出版社,1988:59-92. 
[2] 韩吉人.论职工教育的特点[G]//中国职工教育研究会.职工教育研究论文集.北京:人民教育出版社,1985:90-99. """
    ref.add_text(h)
    ref.render_template()

    ack = Acknowledgments()
    f = """肾衰竭（Kidney failure）是一种终末期的肾脏疾病，此时肾脏的功能会低于其正常水平的15%。由于透析会严重影响患者的生活质量，肾移植一直是治疗肾衰竭的理想方式。但肾脏供体一直处于短缺状态，移植需等待时间约为5-10年。近日，据一篇发表于《美国移植杂志》的文章，阿拉巴马大学伯明翰分校的科学家首次成功将基因编辑猪的肾脏成功移植给一名脑死亡的人类接受者。
研究使用的供体猪的10个关键基因经过了基因编辑，其肾脏的功能更适合人体，且植入人体后引发的免疫排斥反应更轻微。研究人员首先对异种供体和接受者进行交叉配型测试。经配对后，研究人员将肾脏移植入脑死亡接受者的肾脏对应的解剖位置，与肾动脉、肾静脉和输尿管相连接。移植后，他们为患者进行了常规的免疫抑制治疗。目前，肾脏已在患者体内正常工作77小时。该研究按照1期临床试验标准进行，完全按人类供体器官的移植标准实施。研究显示，异种移植的发展在未来或可缓解世界范围器官供应压力。"""
    ack.add_text(f)
    ack.render_template()

    cha = ChangeRecord()
    g = """在无线传能技术中，非辐射无线传能（即电磁感应充电）可以高效传输能量，但有效传输距离被限制在收发器尺寸的几倍之内；而辐射无线传能（如无线电、激光）虽然可以远距离传输能量，但需要复杂的控制机制来跟踪移动的能量接收器。
近日，同济大学电子与信息工程学院的研究团队通过理论和实验证明，"""
    cha.add_text(g)
    cha.render_template()

    apd = Appendixes()
    apd.add_appendix("附录A")
    apd.add_text([Text("啊哈哈哈鸡汤来喽")])
    apd.add_appendix("附录B")
    apd.add_text([Text("直接来吧")])
    apd.render_template()

    #DM.update_toc()
    doc.save("out.docx")