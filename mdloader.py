import markdown
from bs4 import BeautifulSoup, Comment
import logging
import re
from functools import reduce
import os
import docx
import bibtexparser
from bibtexparser.bparser import BibTexParser
from typing import Dict, List

from mdext import MDExt
import dut_paper as word
import dut_paper_translation as transword

file_dir = ""
debug = False


# 检查

def assert_warning(e: bool, s: str):
    if not e:
        logging.warning(s)
    return e


def assert_error(e: bool, s: str):
    if not e:
        logging.error(s)
        exit(-1)
    return e


def log_error(s: str):
    return assert_error(False, s)


def log_warning(s: str):
    return assert_warning(False, s)


# 处理文本

def rbk(text: str):  # remove_blank
    # 删除换行符
    text = text.replace("\n", " ")
    text = text.replace("\r", "")

    cn_char = u'[\u4e00-\u9fa5。，：《》、（）“”‘’]'
    # 中文字符后空格
    should_replace_list = re.compile(
        cn_char + u' +').findall(text)
    # 中文字符前空格
    should_replace_list += re.compile(
        u' +' + cn_char).findall(text)
    # 删除空格
    for i in should_replace_list:
        if i == u' ':
            continue
        new_i = i.strip()
        text = text.replace(i, new_i)
    return text


def raw_text(runs):
    strs = [i["text"] for i in runs]
    return reduce(lambda x, y: x+y, strs)


def assemble_ps(ps):
    strs = []
    for (_, runs) in ps:
        strs.append(raw_text(runs))
    return reduce(lambda x, y: x+"\n"+y, strs)


def split_title(title):
    assert_error(len(title.split(':')) >= 2, "应该有别名或者标题: " + title)
    ali = title.split(':')[0]
    title = rbk(title[len(ali)+1:].strip())
    return ali, title


def ref_items_list_unfold(ref_items_list: list):
    unfold_ref_items = {}
    for ref_items in ref_items_list:
        for ali in ref_items:
            assert_warning(ali not in unfold_ref_items,
                           "任意类型的引用别名应该唯一: " + ali)
            unfold_ref_items[ali] = ref_items[ali]
    return unfold_ref_items


# 数据类型

class RefItem:
    IMG = "img"
    TABLE = "Table"
    MATH = "math"
    LITER = "literature"

    def __init__(self, index, text: str, type: str):
        self.index = str(index)
        self.text = text
        self.type = type


# 每个论文模块

class PaperPart:
    def __init__(self):
        self.contents = []
        self.block: word.Component = None

    # 获取内容

    def load_contents(self, soup: BeautifulSoup): pass

    def _get_content_until(self, cur, until, ollevel=4):
        conts = []
        head_counter = [0]
        while cur != until:
            if cur.name == None:
                cur = cur.next_sibling
                continue
            if cur.name[0] == "h":  # h1 h2 h3
                head_counter, pair = self._process_headline(head_counter,
                                                            cur.name, cur.text)
                conts.append(pair)
            elif cur.name == "p":
                conts += self._process_ps(cur)
            elif cur.name == "table":
                table_name = raw_text(conts[-1][1])
                conts = conts[:-1]
                conts.append(self._process_table(table_name, cur))
            elif cur.name == "ol":
                conts += self._process_ol(cur, ollevel)
            elif cur.name == "math":
                math_title = raw_text(conts[-1][1])
                conts = conts[:-1]
                conts.append(self._process_math(math_title, cur))
            else:
                log_error("这是啥？" + cur.prettify())
            cur = cur.next_sibling
        return conts

    def _get_content_from(self, cur, ollevel=4):
        return self._get_content_until(cur, None, ollevel)

    # 处理标签

    def _process_headline(self, head_counter: List[int], h_label: str, headline: str):
        level = int(h_label[1:])
        assert_warning(1 <= level and level <= len(head_counter)+1,
                       "标题层级应该递进" + headline)
        if level == len(head_counter) + 1:  # new sub part
            head_counter.append(1)
        elif 1 <= level and level <= len(head_counter):  # new part
            head_counter[level-1] += 1
            head_counter = head_counter[:level]
        else:
            log_error("错误的标题编号")

        index = str(head_counter[0])
        for i in range(1, len(head_counter)):
            index += "." + str(head_counter[i])

        headline = headline.strip()
        assert_warning(headline[:len(index)] == index,
                       "没有编号或者编号错误: {} {}".format(h_label, headline))
        assert_warning(headline[len(index)] == " " and
                       headline[len(index)+1] != " ",
                       "MD 中编号后应该有一个空格: {} {}".format(h_label, headline))
        headline = headline[:len(index)] + "  " + rbk(headline[len(index)+1:])

        return head_counter, (h_label, headline)

    def _process_ps(self, p, ollevel=4):
        ps = []
        data = []
        for i in p.children:
            if i.name == None:
                if i.text == "\n":
                    continue
                data.append({"type": "text", "text": rbk(i.text)})
            elif i.name == "strong":
                assert_warning(len(i.contents) == 1, "只允许粗斜体，不允许复杂嵌套")
                if i.contents[0].name == "em":
                    data.append({"type": "strong-em", "text": rbk(i.text)})
                else:
                    data.append({"type": "strong", "text": rbk(i.text)})
            elif i.name == "em":
                data.append({"type": "em", "text": rbk(i.text)})
            elif i.name == "math-inline":
                data.append({"type": "math-inline", "text": i.text})
            elif i.name == "ref":
                data.append({"type": "ref", "text": rbk(i.text.strip())})
            else:  # 需要分段
                if data:
                    ps.append(("p", data))
                    data = []
                if i.name == "br":  # 分段
                    pass
                elif i.name == "img":  # 图片
                    ps.append(self._process_img(i))
                elif i.name == "ol":
                    ps += self._process_ol(i, ollevel)
                else:
                    log_error("缺了什么？" + str(i))
        if data:
            ps.append(("p", data))
        return ps

    def _process_img(self, img):
        global file_dir
        img_path = os.path.join(file_dir, img["src"])
        ali, title = split_title(img["alt"])
        return ("img", {"alias": ali,
                        "title": title,
                        "src": img_path})

    def _process_table(self, title, table):
        data = []
        # 表头，有上实线
        data.append(word.Row([rbk(i.text) for i in table.find("thead").find_all("th")],
                             top_border=True))
        has_border = True  # 表身第一行有上实线
        for tr in table.find("tbody").find_all("tr"):
            row = [rbk(i.text) for i in tr.find_all("td")]  # get all text
            row = list(map(lambda x: None if x == '' else x,
                           row))  # replace '' with None
            if has_border:
                data.append(word.Row(row, top_border=True))
                has_border = False
            else:
                is_border = True
                for i in row:
                    if i == None:
                        is_border = False
                        break
                    for j in i:
                        if j != '-':
                            is_border = False
                            break
                if is_border:
                    has_border = True  # 自定义的实线，下一行数据有上实线
                else:
                    data.append(word.Row(row))

        ali, title = split_title(title)
        return ("table", {"alias": ali,
                          "title": title,
                          "data": data})

    def _process_lis(self, li, level):
        if (li.contents[0].text == "\n"):  # <p>
            conts = self._get_content_from(li.contents[0], level+1)
        else:  # text
            conts = self._process_ps(li, level+1)
        conts[0] = ("fh" + str(level), conts[0][1])
        return conts

    def _process_ol(self, ol, level):
        assert_error(level <= 5, "层次至多两层")
        datas = [self._process_lis(i, level)
                 for i in ol.find_all("li", recursive=False)]
        # make index
        for i in range(len(datas)):
            li_data = datas[i][0]
            if level == 4:
                li_data[1].insert(
                    0, {"type": "text", "text": "（{}） ".format(i+1)})
            else:
                assert_warning(i < 20, "层次二不能超过 20 项")
                li_data[1].insert(
                    0, {"type": "text", "text": "{} ".format(chr(i+0x2460))})  # get ①②..⑳
        data = reduce(lambda x, y: x + y, datas)
        return data

    def _process_math(self, title, math):
        return ("math", {"alias": title,
                         "title": "",
                         "text": math.text})

    # 处理

    def check(self): pass

    def compile(self): pass

    def _get_ref_items(self, conts, index_prefix: str = "") -> Dict[str, RefItem]:
        def get_index(index_prefix: str, chapter_cnt: int, item_cnt: int):
            if index_prefix == "":
                return "{}.{}".format(chapter_cnt, item_cnt)
            else:
                return "{}{}".format(index_prefix, item_cnt)

        def img_index() -> str:
            return get_index(index_prefix, chapter_cnt, img_cnt)

        def table_index() -> str:
            return get_index(index_prefix, chapter_cnt, img_cnt)

        def math_index() -> str:
            return get_index(index_prefix, chapter_cnt, img_cnt)

        ref_items = {}
        chapter_cnt = 0

        for name, cont in conts:
            if name == "h1":
                chapter_cnt += 1
                img_cnt = 0
                table_cnt = 0
                formula_cnt = 0
            elif name in ["img", "table", "math"]:
                ali = cont['alias']
                assert_warning(ali not in ref_items, "有重复别名" + ali)
                if name == "img":
                    img_cnt += 1
                    ref_items[ali] = RefItem(
                        img_index(), "图" + img_index(), RefItem.IMG)
                    cont['title'] = "图{}  {}".format(
                        img_index(), cont['title'])
                elif name == "table":
                    table_cnt += 1
                    ref_items[ali] = RefItem(
                        table_index(), "表" + table_index(), RefItem.TABLE)
                    cont['title'] = "表{}  {}".format(
                        table_index(), cont['title'])
                elif name == "math":
                    formula_cnt += 1
                    ref_items[ali] = RefItem(
                        math_index(), "式" + math_index(), RefItem.MATH)
                    cont['title'] = "（{}）".format(
                        table_index())
        return ref_items

    def get_ref_items(self):
        return self._get_ref_items(self.contents)

    def link_ref(self, ref_items: Dict[str, RefItem], liter_cnt: int) -> int:
        for name, cont in self.contents:
            if name not in ["p", "fh4", "fh5"]:
                continue
            is_text = False
            for run in cont:
                if run["type"] != "ref":
                    if run["type"] == "text" and run["text"].endswith("文献"):
                        is_text = True
                    continue
                ali = run["text"]
                if ali.find(",") == -1:
                    if ali not in ref_items:
                        ref_items[ali] = RefItem(
                            liter_cnt+1, "", RefItem.LITER)
                        liter_cnt += 1
                    ref_item = ref_items[ali]
                    if ref_item.type == RefItem.LITER:
                        run["text"] = "[{}]".format(
                            ref_item.index)
                        if is_text:
                            run["type"] = "text"
                        else:
                            run["type"] = "ref"
                    else:
                        run["type"] = "text"
                        run["text"] = ref_item.text
                else:
                    alis = ali.split(",")
                    for ali in alis:
                        if ali not in ref_items:
                            ref_items[ali] = RefItem(
                                liter_cnt+1, "", RefItem.LITER)
                            liter_cnt += 1
                    index_list = []
                    for ali in alis:
                        assert_error(ref_items[ali].type == RefItem.LITER,
                                     "只有参考文献可以一次引用多个: "+str(alis))
                        index_list.append(int(ref_items[ali].index))
                    index_list.sort()
                    index_list = [[x, x] for x in index_list]
                    sort_list = index_list[:1]
                    for index_pair in index_list[1:]:
                        if sort_list[-1][1]+1 == index_pair[0]:
                            sort_list[-1][1] = index_pair[1]
                        else:
                            sort_list.append(index_pair)
                    sort_list = [str(x[0]) if x[0] == x[1]
                                 else "{}-{}".format(x[0], x[1])
                                 for x in sort_list]
                    run["text"] = "[{}]".format(
                        reduce(lambda x, y: x+","+y, sort_list))
                    if is_text:
                        run["type"] = "text"
                    else:
                        run["type"] = "ref"
                is_text = False
        return liter_cnt

    # 渲染

    def _block_load_body(self, conts=None):
        if conts == None:
            conts = self.contents
        for (name, cont) in conts:
            if name == "h1":
                self.block.add_chapter(cont)
            elif name == "h2":
                self.block.add_section(cont)
            elif name == "h3":
                self.block.add_subsection(cont)
            elif name in ["p", "fh4", "fh5"]:
                if not debug:
                    para = word.Text()
                else:
                    para = word.Text(name)
                for run in cont:
                    if run["type"] == "text":
                        para.add_run(word.Run(run["text"], word.Run.Normal))
                    elif run["type"] == "strong":
                        para.add_run(word.Run(run["text"], word.Run.Bold))
                    elif run["type"] == "em":
                        para.add_run(word.Run(run["text"], word.Run.Italics))
                    elif run["type"] == "strong-em":
                        para.add_run(word.Run(run["text"],
                                              word.Run.Italics | word.Run.Bold))
                    elif run["type"] == "math-inline":
                        para.add_run(word.Run(run["text"], word.Run.Formula))
                    elif run["type"] == "ref":
                        para.add_run(
                            word.Run(run["text"], word.Run.Superscript))
                    else:
                        print("还没实现now", name)
                self.block.add_text([para])
            elif name == "img":
                img = word.Image([word.ImageData(cont["src"], cont["title"])])
                self.block.add_text([img])
            elif name == "table":
                table = word.Table(cont['title'], cont['data'])
                self.block.add_text([table])
            elif name == "math":
                formula = word.Formula(cont['title'], cont['text'])
                self.block.add_text([formula])
            else:
                print("还没实现now", name)

    def _block_load_contents(self):
        self._block_load_body()

    def render(self):
        self._block_load_contents()
        self.block.render_template()


class MetaPart(PaperPart):
    def load_contents(self, soup: BeautifulSoup):
        mete_h1 = soup.find("h1")

        data_table = mete_h1.find_next_sibling("table").find("tbody")
        data_lines = data_table.find_all("tr")
        data_pairs = [list(map(lambda x: rbk(x.text), i.find_all("td")))
                      for i in data_lines]
        data_dict = dict(data_pairs)

        self.title_zh_CN = rbk(mete_h1.text)
        self.title_en = rbk(mete_h1.find_next_sibling("h2").text)
        self.school = data_dict["学院（系）"]
        self.major = data_dict["专业"]
        self.name = data_dict["学生姓名"]
        self.number = data_dict["学号"]
        self.teacher = data_dict["指导教师"]
        self.auditor = data_dict["评阅教师"]
        self.finish_date = data_dict["完成日期"]

    def _block_load_contents(self):
        self.block = word.Metadata()

        self.block.title_zh_CN = self.title_zh_CN
        self.block.title_en = self.title_en
        self.block.school = self.school
        self.block.major = self.major
        self.block.name = self.name
        self.block.number = self.number
        self.block.teacher = self.teacher
        self.block.auditor = self.auditor
        self.block.finish_date = self.finish_date


class AbsPart(PaperPart):
    def load_contents(self, soup: BeautifulSoup):
        # 摘要
        abs_cn_h1 = soup.find("h1", string=re.compile("摘要"))
        abs_cn_ul = abs_cn_h1.find_next_sibling("ul")
        conts_cn = self._get_content_until(abs_cn_h1.next_sibling, abs_cn_ul)
        assert_warning(conts_cn[-1] == ("p", [{"type": "text", "text": "关键词："}]),
                       '摘要应该以"关键词："后接关键词列表结尾')
        self.conts_zh_CN = conts_cn[:-1]
        self.keywords_zh_CN = [rbk(i.text)
                               for i in abs_cn_h1.find_next_sibling("ul").find_all("li")]
        self.title_zh_CN = ""

        # Abstract
        abs_en_h1 = soup.find("h1", string=re.compile("Abstract"))
        abs_en_ul = abs_en_h1.find_next_sibling("ul")
        conts_en = self._get_content_until(abs_en_h1.next_sibling, abs_en_ul)
        assert_warning(conts_en[-1] == ("p", [{"type": "text", "text": "Key Words:"}]),
                       'Abstract应该以"Key Words:"后接关键词列表结尾')
        self.conts_en = conts_en[:-1]
        self.keywords_en = [rbk(i.text)
                            for i in abs_en_h1.find_next_sibling("ul").find_all("li")]
        self.title_en = ""

    def _block_load_contents(self):
        self.block = word.Abstract()
        self.block.set_title(self.title_zh_CN,
                             self.title_en)
        self.block.add_text(assemble_ps(self.conts_zh_CN),
                            assemble_ps(self.conts_en))
        self.block.set_keyword(self.keywords_zh_CN,
                               self.keywords_en)


class IntroPart(PaperPart):
    def load_contents(self, soup: BeautifulSoup):
        intro_h1 = soup.find("h1", string=re.compile("引言"))
        conts = self._get_content_until(intro_h1.next_sibling,
                                        soup.find("h1", string=re.compile("正文")))
        self.contents = conts

    def _block_load_contents(self):
        self.block = word.Introduction()
        self.block.add_text(assemble_ps(self.contents))


class MainPart(PaperPart):
    def load_contents(self, soup: BeautifulSoup):
        main_h1 = soup.find("h1", string=re.compile("正文"))
        conts = self._get_content_until(main_h1.next_sibling,
                                        soup.find("h1", string=re.compile("结论")))
        self.contents = conts

    def _block_load_contents(self):
        self.block = word.MainContent()
        self._block_load_body()


class ConcPart(PaperPart):
    def load_contents(self, soup: BeautifulSoup):
        conclusion_h1 = soup.find("h1", string=re.compile("结论"))
        if conclusion_h1 == None:
            conclusion_h1 = soup.find("h1", string=re.compile("设计总结"))
        assert_error(conclusion_h1 != None, "应该有结论或设计总结")
        conts = self._get_content_until(conclusion_h1.next_sibling,
                                        soup.find("h1", string=re.compile("参考文献")))
        self.contents = conts
        headline = rbk(conclusion_h1.text)
        assert_warning(headline in ["结论", "设计总结"],
                       "结论部分的标题应该是结论/设计总结: "+headline)
        if headline == "结论":
            headline = "结    论"
        self.headline = headline

    def _block_load_contents(self):
        self.block = word.Conclusion()
        self.block.add_text(assemble_ps(self.contents))

    def render(self):
        self._block_load_contents()
        self.block.render_template(self.headline)


class RefPart(PaperPart):
    def __init__(self):
        super().__init__()
        self.ref_map: Dict[str, str] = {}
        self.ref_list: List[str] = []

    def load_contents(self, soup: BeautifulSoup):
        reference_h1 = soup.find("h1", string=re.compile("参考文献"))
        until_h1 = until_h1 = soup.find("h1", string=re.compile("附录"))
        if until_h1 == None:
            until_h1 = soup.find("h1", string=re.compile("修改记录"))

        self.bib_path = ""
        refs: List[str] = []

        cur = reference_h1.next_sibling
        while cur != until_h1:
            if cur.name != "p":
                cur = cur.next_sibling
                continue
            for i in cur.children:
                if i.name == "code":
                    text = i.text.split("\n")
                    if text[0] == "literature":
                        refs += text[1:]
                    elif text[0] == "bib":
                        bib_path = os.path.join(file_dir, text[1])
                        self.bib_path = bib_path
                    else:
                        log_error("这啥? " + i)
            cur = cur.next_sibling

        for ref_item in refs:
            pos = ref_item.find("]")
            assert_error(ref_item[0] == "[" and pos != -1,
                         "参考文献条目应该以 `[索引]` 开头: " + ref_item)
            ref = ref_item[1: pos]
            item = ref_item[pos+1:].strip()
            assert_warning(ref not in self.ref_map,
                           "参考文献索引不能重复: " + ref_item)
            self.ref_map[ref] = item

    def _block_load_contents(self):
        self.block = word.References()
        self._block_load_body()

    def _ref_get_author(self, data: Dict[str, str]) -> List[str]:
        if data["langid"] == "english":
            names = data["author"].split("and")
            authors = []
            for full_name in names:
                full_name = full_name.split(',')
                last_name = full_name[0].strip()
                name = full_name[1].strip().split(" ")
                name = [x[0] for x in name]
                name = reduce(lambda x, y: x+" "+y, name)
                sort_name = "{} {}".format(last_name, name)
                authors.append(sort_name)
            if len(authors) > 3:
                authors = authors[:3]
                authors.append("et al")
            author = reduce(lambda x, y: x+", "+y, authors)
        elif data["langid"] == "chinese":
            names = data["author"].split("and")
            authors = []
            for full_name in names:
                full_name = full_name.split(',')
                last_name = full_name[0].strip()
                name = full_name[1].strip()
                sort_name = "{}{}".format(last_name, name)
                authors.append(sort_name)
            if len(authors) > 3:
                authors = authors[:3]
                authors.append("等")
            author = reduce(lambda x, y: x+", "+y, authors)
        else:
            log_error("没做"+str(data))
        return author

    def _ref_get_entrytype(self, data: Dict[str, str]) -> str:
        type_map = {"book": "M",
                    "inproceedings": "C",
                    "": "G",
                    "": "N",
                    "article": "J",
                    "phdthesis": "D",
                    "techreport": "R",
                    "misc": "S",
                    "patent": "P",
                    "": "DB",
                    "": "CP",
                    "": "EB",
                    }
        return type_map[data["ENTRYTYPE"]]

    def _ref_get_back(self, data: Dict[str, str]) -> str:
        back = ""
        if "address" in data and "publisher" in data:
            address = data["address"].replace("{", "").replace("}", "")
            publisher = data["publisher"].replace("{", "").replace("}", "")
            back = "{}: {}, ".format(address, publisher)
        return back

    def _ref_GB_T_7714_2005(self, data: Dict[str, str]) -> str:
        assert_error("langid" in data, "参考文献应该有语言信息: "+str(data))
        langid = data["langid"]
        author = self._ref_get_author(data)
        title = data["title"].replace("{", "").replace("}", "")
        entrytype = self._ref_get_entrytype(data)
        year = data["year"].replace("{", "").replace("}", "")
        back = self._ref_get_back(data)

        if langid == "english":
            ref_item = "{}. {} [{}]. {}{}.".format(
                author, title, entrytype, back, year)
        elif langid == "chinese":
            ref_item = "{}. {}[{}]. {}{}.".format(
                author, title, entrytype, back, year)
        else:
            log_error("没做"+str(data))
        return ref_item

    def _load_bib(self) -> Dict[str, str]:
        if self.bib_path == "":
            return {}
        with open(self.bib_path) as bibtex_file:
            parser = BibTexParser(common_strings=True)
            bib_database = bibtexparser.load(bibtex_file, parser=parser)
        ref_map = {}
        for item in bib_database.entries:
            ref_map["@"+item["ID"]] = self._ref_GB_T_7714_2005(item)
        return ref_map

    def compile(self):
        ref_map = self._load_bib()
        for ref in ref_map:
            assert_warning(ref not in self.ref_map,
                           "参考文献索引不能重复: " + ref)
            self.ref_map[ref] = ref_map[ref]

    def filt_ref(self, ref_items: Dict[str, RefItem]):
        ali_list = [(int(ref_items[ali].index), ali)
                    for ali in ref_items
                    if ref_items[ali].type == RefItem.LITER]
        ali_list.sort()
        self.ref_list = []
        for index, ali in ali_list:
            assert_warning(ali in self.ref_map,
                           "引用的文献应该在参考文献中出现: " + ali +
                           " BibTeX_path: '" + self.bib_path + "'")
            if ali in self.ref_map:
                self.ref_list.append(
                    "[{}] {}".format(index, self.ref_map[ali]))
        self.contents = [("p", [{"type": "text", "text": text}])
                         for text in self.ref_list]


class AppenPart(PaperPart):
    class AppenOne:
        def __init__(self, title: str, conts):
            self.title = title
            self.contents = conts

    def __init__(self):
        super().__init__()
        self.appens: List[self.AppenOne] = []

    def load_contents(self, soup: BeautifulSoup):
        appendix_h1s = soup.find_all("h1", string=re.compile("附录"))
        appendix_h1s.append(soup.find("h1", string=re.compile("修改记录")))
        appens = []
        for i in range(0, len(appendix_h1s)-1):
            conts = self._get_content_until(appendix_h1s[i].next_sibling,
                                            appendix_h1s[i+1])
            title = self._process_title(appendix_h1s[i].text, i)
            appens.append(self.AppenOne(title, conts))
        self.appens = appens

    def _block_load_contents(self):
        self.block = word.Appendixes()
        for appen in self.appens:
            self.block.add_appendix(appen.title)
            self._block_load_body(appen.contents)

    def _process_title(self, title: str, index: int):
        assert_warning(title[:2] == "附录", "附录应该以附录和编号开头: " + title)
        if title[2] == ' ':
            title = title[:2] + title[3:]
        assert_warning(title[2] == chr(ord("A") + index),
                       "附录应该以大写字母顺序编号: " + title)
        assert_warning(title[3] == " " and title[4] != " ",
                       "MD 中附录编号后应该有一个空格: " + title)
        title = title[:3] + "  " + rbk(title[4:].strip())
        return title

    def get_ref_items(self):
        ref_items_list = [self._get_ref_items(appen.contents, appen.title[2])
                          for appen in self.appens]
        return ref_items_list_unfold(ref_items_list)


class RecordPart(PaperPart):
    def load_contents(self, soup: BeautifulSoup):
        mod_record_h1 = soup.find("h1", string=re.compile("修改记录"))
        conts = self._get_content_until(mod_record_h1.next_sibling,
                                        soup.find("h1", string=re.compile("致谢")))
        self.contents = conts

    def _block_load_contents(self):
        self.block = word.ChangeRecord()
        self._block_load_body()


class ThanksPart(PaperPart):
    def load_contents(self, soup: BeautifulSoup):
        thanks_h1 = soup.find("h1", string=re.compile("致谢"))
        self.contents = self._get_content_from(thanks_h1.next_sibling)

    def _block_load_contents(self):
        self.block = word.Acknowledgments()
        self.block.add_text(assemble_ps(self.contents))


class TransMetaPart(PaperPart):
    def load_contents(self, soup: BeautifulSoup):
        mete_h1 = soup.find("h1")

        # 个人信息
        data_table = mete_h1.find_next_sibling("table").find("tbody")
        data_lines = data_table.find_all("tr")
        data_pairs = [list(map(lambda x: rbk(x.text), i.find_all("td")))
                      for i in data_lines]
        data_dict = dict(data_pairs)

        self.title_zh_CN = rbk(mete_h1.text)
        self.title_en = rbk(mete_h1.find_next_sibling("h2").text)
        self.school = data_dict["学部（院）"]
        self.major = data_dict["专业"]
        self.name = data_dict["学生姓名"]
        self.number = data_dict["学号"]
        self.teacher = data_dict["指导教师"]
        self.finish_date = data_dict["完成日期"]

        # 外文作者信息
        data_table = mete_h1.find_next_siblings("table")[1].find("tbody")
        data_lines = data_table.find_all("tr")
        data_pairs = [list(map(lambda x: rbk(x.text), i.find_all("td")))
                      for i in data_lines]
        data_dict = dict(data_pairs)

        self.author = data_dict["author"]
        self.organization = data_dict["工作单位"]

    def _block_load_contents(self):
        self.block = transword.TranslationMetadata()

        self.block.title_zh_CN = self.title_zh_CN
        self.block.title_en = self.title_en
        self.block.school = self.school
        self.block.major = self.major
        self.block.name = self.name
        self.block.number = self.number
        self.block.teacher = self.teacher
        self.block.finish_date = self.finish_date


class TransAbsPart(PaperPart):
    def load_contents(self, soup: BeautifulSoup):
        # 摘要
        abs_cn_h1 = soup.find("h1", string=re.compile("摘要"))
        if abs_cn_h1 == None:
            self.conts_zh_CN = None
            return
        abs_cn_ul = abs_cn_h1.find_next_sibling("ul")
        conts_cn = self._get_content_until(abs_cn_h1.next_sibling, abs_cn_ul)
        assert_warning(conts_cn[-1] == ("p", [{"type": "text", "text": "关键词："}]),
                       '摘要应该以"关键词："后接关键词列表结尾')
        self.conts_zh_CN = conts_cn[:-1]
        self.keywords_zh_CN = [rbk(i.text)
                               for i in abs_cn_h1.find_next_sibling("ul").find_all("li")]
        self.title_zh_CN = ""
        self.author = ""
        self.organization = ""

    def _block_load_contents(self):
        self.block = transword.TranslationAbstract(self.title_zh_CN,
                                                   self.author,
                                                   self.organization)
        self._block_load_body(self.conts_zh_CN)
        self.block.add_keywords(self.keywords_zh_CN)


class TransMainPart(PaperPart):
    def load_contents(self, soup: BeautifulSoup):
        main_h1 = soup.find("h1", string=re.compile("正文"))
        conts = self._get_content_from(main_h1.next_sibling)
        self.contents = conts

    def _link_ref(self) -> int:
        for name, cont in self.contents:
            if name not in ["p", "fh4", "fh5"]:
                continue
            is_text = False
            for run in cont:
                if run["type"] != "ref":
                    if run["type"] == "text" and run["text"].endswith("文献"):
                        is_text = True
                    continue

                if is_text:
                    run["type"] = "text"
                else:
                    run["type"] = "ref"
                run["text"] = "[{}]".format(run["text"])

                is_text = False

    def compile(self):
        self._link_ref()

        self.contents.insert(-2, ("p", []))

    def _block_load_contents(self):
        self.block = transword.TranslationMainContent()
        self._block_load_body()


class Paper:
    def __init__(self):
        self.parts: list[PaperPart]
        self.ref_items: Dict[str, Dict[str, str]] = {}

    def load_md(self, md_path: str):
        with open(md_path, "r") as f:
            md_file = f.read()
        global file_dir
        file_dir = os.path.dirname(md_path)
        md_html = markdown.markdown(md_file,
                                    tab_length=3,
                                    extensions=['markdown.extensions.tables',
                                                MDExt()])
        self.soup = BeautifulSoup(md_html, 'html.parser')
        for i in self.soup(text=lambda text: isinstance(text, Comment)):
            i.extract()  # 删除 html 注释

        if debug:
            with open("out.html", "w") as f:
                f.write(self.soup.prettify())

    def load_contents(self):
        for part in self.parts:
            part.load_contents(self.soup)

    def compile(self):
        for part in self.parts:
            part.compile()

    def render(self, doc_path: str, out_path: str):
        doc = docx.Document(doc_path)
        word.DM.set_doc(doc)

        for part in self.parts:
            part.render()

        word.DM.update_toc()
        doc.save(out_path)


class GraduationPaper(Paper):
    def __init__(self):
        super().__init__()
        self.meta = MetaPart()
        self.abs = AbsPart()
        self.intro = IntroPart()
        self.main = MainPart()
        self.conc = ConcPart()
        self.ref = RefPart()
        self.appen = AppenPart()
        self.record = RecordPart()
        self.thanks = ThanksPart()

        self.parts: List[PaperPart] = [
            self.meta,
            self.abs,
            self.intro,
            self.main,
            self.conc,
            self.ref,
            self.appen,
            self.record,
            self.thanks
        ]

    def compile(self):
        super().compile()

        self.abs.title_zh_CN = self.meta.title_zh_CN
        self.abs.title_en = self.meta.title_en

        ref_items_list = [
            self.main.get_ref_items(),
            self.appen.get_ref_items()
        ]
        self.ref_items = ref_items_list_unfold(ref_items_list)
        liter_cnt = 0
        for part in self.parts:
            liter_cnt = part.link_ref(self.ref_items, liter_cnt)
        self.ref.filt_ref(self.ref_items)


class TranslationPaper(Paper):
    def __init__(self):
        super().__init__()
        self.meta = TransMetaPart()
        self.abs = TransAbsPart()
        self.main = TransMainPart()

        self.parts: List[PaperPart] = [
            self.meta,
            self.abs,
            self.main
        ]

    def compile(self):
        super().compile()

        self.abs.author = self.meta.author
        self.abs.organization = self.meta.organization
        self.abs.title_zh_CN = self.meta.title_zh_CN


if __name__ == "__main__":
    paper = GraduationPaper()
    paper.load_md("论文模板.md")
    paper.load_contents()
    paper.compile()

    paper.render("毕业设计（论文）模板-docx.docx", "out.docx")

'''
("h1", "something")
("h2", "something")
("h3", "something")

("fh4", like p)
("fh5", like p)

("p", [("text",      "something"),
       ("strong",    "something"),
       ("strong-em", "something"),
       ("em",        "something")])
("img",     (title, src))
("table",   (title, [Row]))
("formula", (title, "somthing"))
'''
